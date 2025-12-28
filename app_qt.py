import os
import re
import shutil
import subprocess
import sys
import tempfile
import traceback
import json
import zipfile
import xml.etree.ElementTree as ET
import hashlib
from datetime import datetime

from PyQt6 import QtCore, QtGui, QtWidgets
from docx import Document
from docx.shared import Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


def is_windows():
    return sys.platform.startswith("win")


def get_user_data_dir() -> str:
    base = os.environ.get("LOCALAPPDATA") or os.path.expanduser("~")
    path = os.path.join(base, "TuanziConverter")
    os.makedirs(path, exist_ok=True)
    return path


def find_pandoc():
    local = os.path.join(os.path.dirname(__file__), "bin", "pandoc.exe")
    if os.path.isfile(local):
        return local
    return shutil.which("pandoc")


def merge_vertical_text(text: str) -> str:
    lines = text.splitlines()
    out = []
    buffer = []
    short_run = 0
    in_code = False

    def is_short_line(value: str) -> bool:
        if len(value) <= 2:
            return True
        if len(value) <= 4 and not re.search(r"\s", value):
            return bool(re.search(r"[\u4e00-\u9fff\u3000-\u303f]", value))
        return False

    def flush_buffer():
        nonlocal short_run
        if not buffer:
            return
        if short_run >= 3:
            merged = "".join(buffer).strip()
            if merged:
                out.append(merged)
        else:
            out.extend(buffer)
        buffer.clear()
        short_run = 0

    for line in lines:
        stripped = line.strip()
        if stripped.startswith("```"):
            flush_buffer()
            in_code = not in_code
            out.append(line)
            continue
        if in_code:
            out.append(line)
            continue
        if not stripped:
            flush_buffer()
            out.append(line)
            continue
        if is_short_line(stripped):
            buffer.append(stripped)
            short_run += 1
            continue
        flush_buffer()
        out.append(line)
    flush_buffer()
    return "\n".join(out)


def ensure_closed_code_blocks(text: str) -> str:
    lines = text.splitlines()
    out = []
    in_code = False
    for line in lines:
        stripped = line.strip()
        if stripped.startswith("```"):
            in_code = not in_code
        out.append(line)
    if in_code:
        out.append("```")
    return "\n".join(out)


def normalize_ai_headings(text: str) -> str:
    lines = text.splitlines()
    out = []
    in_math = False
    in_code = False
    for line in lines:
        stripped = line.strip()
        if stripped.startswith("```"):
            in_code = not in_code
            out.append(line)
            continue
        if stripped == "$$":
            in_math = not in_math
            out.append(line)
            continue
        if stripped.startswith(r"\begin{"):
            in_math = True
        if in_code or in_math:
            out.append(line)
            if stripped.startswith(r"\end{"):
                in_math = False
            continue
        if not stripped:
            out.append(line)
            continue
        if re.match(r"^#{1,6}\s+", stripped):
            out.append(line)
            continue
        if len(stripped) > 32:
            out.append(line)
            continue
        if re.match(r"^[一二三四五六七八九十]+[、.]\s*\S+", stripped):
            out.append("# " + stripped)
            continue
        if re.match(r"^\d+(?:\.\d+)*[\.、]?\s*\S+", stripped):
            level = stripped.count(".") + 1
            level = min(3, max(1, level))
            out.append("#" * level + " " + stripped)
            continue
        if re.match(r"^\d+[\ufe0f\u20e3]\s*\S+", stripped):
            out.append("# " + stripped)
            continue
        if re.match(r"^第\S{1,3}章\s*\S*", stripped):
            out.append("# " + stripped)
            continue
        out.append(line)
    return "\n".join(out)


def filter_horizontal_rules(text: str, keep_rules: bool) -> str:
    if keep_rules:
        return text
    lines = text.splitlines()
    out = []
    for line in lines:
        stripped = line.strip()
        if stripped and all(ch == "-" for ch in stripped) and len(stripped) >= 3:
            continue
        if stripped and all(ch == "*" for ch in stripped) and len(stripped) >= 3:
            continue
        if stripped and all(ch == "_" for ch in stripped) and len(stripped) >= 3:
            continue
        out.append(line)
    return "\n".join(out)


def strip_links(text: str) -> str:
    # Convert markdown links to plain text and neutralize bare URLs.
    text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)
    text = re.sub(r"<(https?://[^>]+)>", r"\1", text)
    return text


def sanitize_special_marks(text: str) -> str:
    # Remove visible line-break symbols or separators that users copy in.
    text = text.replace("\u21b5", "")
    text = text.replace("\u2028", " ")
    text = text.replace("\u2029", " ")
    return text


def read_text_file(path: str) -> str:
    encodings = ["utf-8-sig", "utf-8", "gb18030", "gbk"]
    for enc in encodings:
        try:
            with open(path, "r", encoding=enc) as f:
                return f.read()
        except UnicodeDecodeError:
            continue
    with open(path, "r", encoding="utf-8", errors="ignore") as f:
        return f.read()


def normalize_math_blocks(text: str) -> str:
    envs = {
        "aligned", "align", "align*", "equation", "equation*", "gather", "gather*",
        "cases", "matrix", "pmatrix", "bmatrix", "vmatrix", "Vmatrix",
    }
    split_envs = {"aligned", "align", "align*", "gather", "gather*"}
    lines = text.splitlines()
    out = []
    in_env = False
    buffer = []
    env_name = None

    def flush():
        nonlocal in_env, env_name
        if not buffer:
            return
        content = "\n".join(buffer)
        if env_name in split_envs:
            inner = content
            inner = re.sub(r"^\\s*\\\\begin\\{[^}]+\\}\\s*", "", inner)
            inner = re.sub(r"\\s*\\\\end\\{[^}]+\\}\\s*$", "", inner)
            parts = [p.strip() for p in inner.split(r"\\\\") if p.strip()]
            for part in parts:
                out.append("")
                out.append("$$")
                out.append(part.replace("#", r"\#"))
                out.append("$$")
                out.append("")
        else:
            out.append("")
            out.append("$$")
            for ln in buffer:
                cleaned = ln.replace("#", r"\#")
                out.append(cleaned)
            out.append("$$")
            out.append("")
        buffer.clear()
        in_env = False
        env_name = None

    for line in lines:
        stripped = line.strip()
        if stripped.startswith(r"\begin{"):
            name = stripped[len(r"\begin{"):].split("}", 1)[0]
            if name in envs:
                in_env = True
                env_name = name
                buffer.append(stripped)
                continue
        if in_env:
            buffer.append(stripped)
            if stripped.startswith(r"\end{") and env_name in stripped:
                flush()
            continue
        out.append(line)
    flush()
    return "\n".join(out)


def normalize_fenced_math(text: str) -> str:
    lines = text.splitlines()
    out = []
    in_math = False
    buffer = []

    def flush():
        if not buffer:
            return
        out.append("")
        out.append("$$")
        for ln in buffer:
            out.append(ln)
        out.append("$$")
        out.append("")
        buffer.clear()

    for line in lines:
        stripped = line.strip()
        if stripped.startswith("```math"):
            in_math = True
            continue
        if in_math and stripped.startswith("```"):
            in_math = False
            flush()
            continue
        if in_math:
            buffer.append(line)
            continue
        out.append(line)
    flush()
    return "\n".join(out)


def strip_standalone_brackets(text: str) -> str:
    lines = text.splitlines()
    out = []
    in_bracket = False
    buffer = []

    def flush_as_math():
        if not buffer:
            return
        out.append("")
        out.append("$$")
        out.extend(buffer)
        out.append("$$")
        out.append("")
        buffer.clear()

    for line in lines:
        stripped = line.strip()
        if stripped in {"[", r"\["}:
            if in_bracket:
                out.append(line)
            else:
                in_bracket = True
                buffer.clear()
            continue
        if stripped in {"]", r"\]"}:
            if in_bracket:
                flush_as_math()
                in_bracket = False
            else:
                continue
            continue
        if in_bracket:
            buffer.append(line)
            continue
        out.append(line)
    if in_bracket:
        out.append("[")
        out.extend(buffer)
    return "\n".join(out)


def normalize_inline_math_parens(text: str) -> str:
    lines = text.splitlines()
    out = []
    in_code = False
    in_math = False
    math_hint = re.compile(r"\\[A-Za-z]+|[_^]")

    def convert_segment(segment: str) -> str:
        def process(s: str) -> str:
            res = []
            i = 0
            while i < len(s):
                if s[i] == "(":
                    depth = 1
                    j = i + 1
                    while j < len(s) and depth > 0:
                        if s[j] == "(":
                            depth += 1
                        elif s[j] == ")":
                            depth -= 1
                        j += 1
                    if depth == 0:
                        inner = s[i + 1 : j - 1]
                        if inner and "$" not in inner and math_hint.search(inner):
                            res.append("$" + inner.strip() + "$")
                        else:
                            res.append("(" + process(inner) + ")")
                        i = j
                        continue
                res.append(s[i])
                i += 1
            return "".join(res)

        return process(segment)

    for line in lines:
        stripped = line.strip()
        if stripped.startswith("```"):
            in_code = not in_code
            out.append(line)
            continue
        if stripped == "$$":
            in_math = not in_math
            out.append(line)
            continue
        if in_code or in_math:
            out.append(line)
            continue
        parts = line.split("`")
        for idx in range(0, len(parts), 2):
            parts[idx] = convert_segment(parts[idx])
        out.append("`".join(parts))
    return "\n".join(out)


def unescape_markdown_from_docx(text: str) -> str:
    # Pandoc escapes markdown symbols when source docx contains literal markdown.
    # We only unescape punctuation, not LaTeX commands.
    text = re.sub(r"\\(?=[#*`_>\\[\\]-])", "", text)
    return text


def update_page_numbers(docx_path: str, enable: bool, start: int, start_page: int):
    if not docx_path or not os.path.isfile(docx_path):
        return
    ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    ns_rel = "http://schemas.openxmlformats.org/package/2006/relationships"
    ns_r = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
    ET.register_namespace("w", ns)
    ET.register_namespace("r", ns_r)

    def has_page_field(para):
        for node in para.iter():
            if node.tag == "{" + ns + "}fldSimple":
                instr = node.get("{" + ns + "}instr", "") or ""
                if "PAGE" in instr or "NUMPAGES" in instr:
                    return True
            if node.tag == "{" + ns + "}instrText":
                text = node.text or ""
                if "PAGE" in text or "NUMPAGES" in text:
                    return True
        return False

    try:
        with zipfile.ZipFile(docx_path, "r") as z:
            names = z.namelist()
            data = z.read("word/document.xml") if "word/document.xml" in names else None
            rels_data = z.read("word/_rels/document.xml.rels") if "word/_rels/document.xml.rels" in names else None
            footers = {n: z.read(n) for n in names if n.startswith("word/footer")}
        if data is None or rels_data is None:
            return
        doc_root = ET.fromstring(data)
        rels_root = ET.fromstring(rels_data)
        rels = rels_root.findall(".//{" + ns_rel + "}Relationship")
        rel_map = {r.get("Id"): r.get("Target") for r in rels}
        used_rids = {r.get("Id") for r in rels}
        max_footer_idx = 0
        for name in footers.keys():
            m = re.search(r"footer(\\d+)\\.xml", name)
            if m:
                max_footer_idx = max(max_footer_idx, int(m.group(1)))

        def strip_page_fields(xml_bytes):
            root = ET.fromstring(xml_bytes)
            parent_map = {c: p for p in root.iter() for c in p}
            for para in list(root.findall(".//{" + ns + "}p")):
                if has_page_field(para):
                    parent = parent_map.get(para)
                    if parent is None:
                        root.remove(para)
                    else:
                        parent.remove(para)
            return ET.tostring(root, encoding="utf-8", xml_declaration=True)

        def to_part_path(target: str):
            if target.startswith("word/"):
                return target
            return "word/" + target

        footer_updates = {name: raw for name, raw in footers.items()}

        for sect in doc_root.findall(".//{" + ns + "}sectPr"):
            pg = sect.find("{" + ns + "}pgNumType")
            if enable:
                if pg is None:
                    pg = ET.SubElement(sect, "{" + ns + "}pgNumType")
                pg.set("{" + ns + "}start", str(max(1, int(start))))
            else:
                if pg is not None:
                    sect.remove(pg)

            default_ref = None
            first_ref = None
            for ref in list(sect.findall("{" + ns + "}footerReference")):
                ref_type = ref.get("{" + ns + "}type")
                if ref_type == "default":
                    default_ref = ref
                elif ref_type == "first":
                    first_ref = ref

            if enable and start_page > 1:
                title_pg = sect.find("{" + ns + "}titlePg")
                if title_pg is None:
                    sect.append(ET.Element("{" + ns + "}titlePg"))
                if first_ref is None:
                    max_footer_idx += 1
                    new_footer = f"footer{max_footer_idx}.xml"
                    new_part = "word/" + new_footer
                    idx = 1
                    while f"rId{idx}" in used_rids:
                        idx += 1
                    new_rel_id = f"rId{idx}"
                    used_rids.add(new_rel_id)
                    rel = ET.Element("{" + ns_rel + "}Relationship")
                    rel.set("Id", new_rel_id)
                    rel.set("Type", "http://schemas.openxmlformats.org/officeDocument/2006/relationships/footer")
                    rel.set("Target", new_footer)
                    rels_root.append(rel)
                    first_ref = ET.Element("{" + ns + "}footerReference")
                    first_ref.set("{" + ns + "}type", "first")
                    first_ref.set("{" + ns_r + "}id", new_rel_id)
                    sect.append(first_ref)
                    if default_ref is not None:
                        def_target = rel_map.get(default_ref.get("{" + ns_r + "}id"))
                        def_part = to_part_path(def_target) if def_target else None
                        if def_part and def_part in footers:
                            footer_updates[new_part] = strip_page_fields(footers[def_part])
                        else:
                            footer_updates[new_part] = (
                                b'<?xml version="1.0" encoding="UTF-8"?>'
                                b'<w:ftr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
                                b'<w:p/></w:ftr>'
                            )
                    else:
                        footer_updates[new_part] = (
                            b'<?xml version="1.0" encoding="UTF-8"?>'
                            b'<w:ftr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
                            b'<w:p/></w:ftr>'
                        )
                else:
                    first_target = rel_map.get(first_ref.get("{" + ns_r + "}id"))
                    first_part = to_part_path(first_target) if first_target else None
                    if first_part and first_part in footers:
                        footer_updates[first_part] = strip_page_fields(footers[first_part])
            else:
                title_pg = sect.find("{" + ns + "}titlePg")
                if title_pg is not None:
                    sect.remove(title_pg)
                if first_ref is not None:
                    sect.remove(first_ref)

        if not enable and footers:
            for name, raw in footers.items():
                footer_updates[name] = strip_page_fields(raw)

        new_doc = ET.tostring(doc_root, encoding="utf-8", xml_declaration=True)
        new_rels = ET.tostring(rels_root, encoding="utf-8", xml_declaration=True)
        tmp_path = docx_path + ".tmp"
        with zipfile.ZipFile(docx_path, "r") as zin, zipfile.ZipFile(tmp_path, "w", compression=zipfile.ZIP_DEFLATED) as zout:
            existing = {i.filename for i in zin.infolist()}
            for item in zin.infolist():
                if item.filename == "word/document.xml":
                    zout.writestr(item.filename, new_doc)
                elif item.filename == "word/_rels/document.xml.rels":
                    zout.writestr(item.filename, new_rels)
                elif item.filename in footer_updates:
                    zout.writestr(item.filename, footer_updates[item.filename])
                else:
                    zout.writestr(item.filename, zin.read(item.filename))
            for name, raw in footer_updates.items():
                if name not in existing:
                    zout.writestr(name, raw)
        os.replace(tmp_path, docx_path)
    except Exception:
        return


def normalize_table_breaks(text: str) -> str:
    lines = text.splitlines()
    out = []
    for line in lines:
        if line.count("|") >= 2:
            cleaned = re.sub(r"<br\\s*/?>", " ", line, flags=re.IGNORECASE)
            cleaned = cleaned.replace("\\\\", " ")
            out.append(cleaned)
        else:
            out.append(line)
    return "\n".join(out)


def normalize_markdown_tables(text: str) -> str:
    lines = text.splitlines()
    out = []
    for line in lines:
        if line.count("|") >= 2:
            cleaned = re.sub(r"<br\\s*/?>", " ", line, flags=re.IGNORECASE)
            cleaned = cleaned.replace("\\\\", " ")
            stripped = cleaned.strip()
            parts = [p.strip() for p in stripped.strip("|").split("|")]
            if re.fullmatch(r"\|?[\s:-]+\|?[\s:-]*", stripped):
                norm = []
                for p in parts:
                    p = p.replace(" ", "")
                    if ":" in p:
                        left = p.startswith(":")
                        right = p.endswith(":")
                        mid = "---"
                        if left and right:
                            norm.append(f":{mid}:")
                        elif left:
                            norm.append(f":{mid}")
                        elif right:
                            norm.append(f"{mid}:")
                        else:
                            norm.append(mid)
                    else:
                        norm.append("---")
                out.append("| " + " | ".join(norm) + " |")
            else:
                out.append("| " + " | ".join(parts) + " |")
        else:
            out.append(line)
    return "\n".join(out)


def remove_standalone_dollars(text: str) -> str:
    lines = text.splitlines()
    return "\n".join([ln for ln in lines if ln.strip() != "$$"])


def extract_tables_with_placeholders(text: str):
    placeholder_prefix = "TABLE_PLACEHOLDER_"
    lines = text.splitlines()
    tables = []
    out = []
    i = 0
    in_code = False
    fullwidth_pipe = "｜"
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()
        if stripped.startswith("```"):
            in_code = not in_code
            out.append(line)
            i += 1
            continue
        if in_code:
            out.append(line)
            i += 1
            continue
        scan_line = line.replace(fullwidth_pipe, "|")
        if scan_line.count("|") >= 2:
            block = []
            block_cleaned = []
            while i < len(lines):
                scan = lines[i].replace(fullwidth_pipe, "|")
                if scan.count("|") < 2:
                    break
                block.append(lines[i])
                cleaned_line = lines[i].replace(fullwidth_pipe, "|")
                cleaned_line = re.sub(r"<br\\s*/?>", " ", cleaned_line, flags=re.IGNORECASE)
                cleaned_line = cleaned_line.replace("\\\\", " ")
                cleaned_line = sanitize_special_marks(cleaned_line)
                block_cleaned.append(cleaned_line)
                i += 1
            rows = []
            has_math = False
            for bl in block:
                cleaned = bl.replace(fullwidth_pipe, "|")
                cleaned = re.sub(r"<br\\s*/?>", " ", cleaned, flags=re.IGNORECASE)
                cleaned = cleaned.replace("\\\\", " ")
                cleaned = sanitize_special_marks(cleaned)
                stripped = cleaned.strip()
                if re.fullmatch(r"\|?[\s:-]+\|?[\s:-]*", stripped):
                    continue
                parts = [p.strip() for p in stripped.strip("|").split("|")]
                parts = [re.sub(r"[\r\n]+", " ", p) for p in parts]
                if not has_math:
                    for p in parts:
                        if re.search(r"\\[A-Za-z]+|[_^]", p):
                            has_math = True
                            break
                rows.append(parts)
            if rows:
                if has_math:
                    out.extend(block_cleaned or block)
                else:
                    tables.append(rows)
                    out.append(f"{placeholder_prefix}{len(tables)}")
            else:
                out.extend(block_cleaned or block)
            continue
        out.append(line)
        i += 1
    return "\n".join(out), tables


def rebuild_tables_in_docx(docx_path: str, tables_data):
    if not tables_data:
        return
    doc = Document(docx_path)
    table_style = None
    try:
        table_style = doc.styles["Table Grid"].name
    except Exception:
        table_style = None
    placeholder_re = re.compile(r"^TABLE_PLACEHOLDER_(\d+)$")
    for paragraph in list(doc.paragraphs):
        text = paragraph.text.strip()
        match = placeholder_re.match(text)
        if not match:
            continue
        idx = int(match.group(1)) - 1
        if idx < 0 or idx >= len(tables_data):
            continue
        table_data = tables_data[idx]
        if not table_data:
            continue
        cols = max(len(r) for r in table_data)
        table = doc.add_table(rows=0, cols=cols)
        if table_style:
            table.style = table_style
        for row in table_data:
            cells = table.add_row().cells
            for idx in range(cols):
                cells[idx].text = row[idx] if idx < len(row) else ""
        paragraph._p.addprevious(table._tbl)
        paragraph._element.getparent().remove(paragraph._element)
    doc.save(docx_path)


def apply_table_grid_style(docx_path: str):
    if not docx_path or not os.path.isfile(docx_path):
        return
    try:
        doc = Document(docx_path)
        style_names = {s.name for s in doc.styles}
        preferred = None
        for name in ("Table Grid", "表格网格", "网格型表格"):
            if name in style_names:
                preferred = name
                break
        for table in doc.tables:
            try:
                if preferred:
                    table.style = preferred
                else:
                    set_table_borders(table)
            except Exception:
                pass
        doc.save(docx_path)
    except Exception:
        return


def set_table_borders(table, size: int = 8, color: str = "000000"):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        tbl.append(tbl_pr)
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), str(size))
        element.set(qn("w:color"), color)
        element.set(qn("w:space"), "0")


def patch_numbering_fonts(docx_path: str, east_asia: str, latin: str):
    if not docx_path or not os.path.isfile(docx_path):
        return
    if not east_asia and not latin:
        return
    ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    ET.register_namespace("w", ns)
    try:
        with zipfile.ZipFile(docx_path, "r") as z:
            if "word/numbering.xml" not in z.namelist():
                return
            data = z.read("word/numbering.xml")
        root = ET.fromstring(data)
        for lvl in root.findall(".//{"+ns+"}lvl"):
            numFmt = lvl.find("{"+ns+"}numFmt")
            if numFmt is None:
                continue
            val = numFmt.get("{"+ns+"}val", "")
            if "decimal" not in val:
                continue
            rPr = lvl.find("{"+ns+"}rPr")
            if rPr is None:
                rPr = ET.SubElement(lvl, "{"+ns+"}rPr")
            rFonts = rPr.find("{"+ns+"}rFonts")
            if rFonts is None:
                rFonts = ET.SubElement(rPr, "{"+ns+"}rFonts")
            if latin:
                rFonts.set("{"+ns+"}ascii", latin)
                rFonts.set("{"+ns+"}hAnsi", latin)
                rFonts.set("{"+ns+"}cs", latin)
            if east_asia:
                rFonts.set("{"+ns+"}eastAsia", east_asia)
        new_data = ET.tostring(root, encoding="utf-8", xml_declaration=True)
        tmp_path = docx_path + ".tmp"
        with zipfile.ZipFile(docx_path, "r") as zin, zipfile.ZipFile(tmp_path, "w", compression=zipfile.ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                if item.filename == "word/numbering.xml":
                    zout.writestr(item, new_data)
                else:
                    zout.writestr(item, zin.read(item.filename))
        os.replace(tmp_path, docx_path)
    except Exception:
        return


def patch_styles_fonts(docx_path: str, east_asia: str, latin: str):
    if not docx_path or not os.path.isfile(docx_path):
        return
    if not east_asia and not latin:
        return
    ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    ET.register_namespace("w", ns)
    try:
        with zipfile.ZipFile(docx_path, "r") as z:
            if "word/styles.xml" not in z.namelist():
                return
            data = z.read("word/styles.xml")
        root = ET.fromstring(data)

        def set_fonts(rFonts):
            if latin:
                rFonts.set("{"+ns+"}ascii", latin)
                rFonts.set("{"+ns+"}hAnsi", latin)
                rFonts.set("{"+ns+"}cs", latin)
                for attr in ("asciiTheme", "hAnsiTheme", "cstheme"):
                    key = "{"+ns+"}"+attr
                    if key in rFonts.attrib:
                        del rFonts.attrib[key]
            if east_asia:
                rFonts.set("{"+ns+"}eastAsia", east_asia)
                key = "{"+ns+"}eastAsiaTheme"
                if key in rFonts.attrib:
                    del rFonts.attrib[key]

        # docDefaults
        rpr_default = root.find(".//{"+ns+"}rPrDefault/{"+ns+"}rPr")
        if rpr_default is not None:
            rFonts = rpr_default.find("{"+ns+"}rFonts")
            if rFonts is None:
                rFonts = ET.SubElement(rpr_default, "{"+ns+"}rFonts")
            set_fonts(rFonts)

        target_ids = {
            "Normal", "BodyText", "BodyText2", "BodyText3", "Title", "Subtitle",
            "Heading1", "Heading2", "Heading3", "Heading4", "Heading5", "Heading6",
            "Heading1Char", "Heading2Char", "Heading3Char", "Heading4Char", "Heading5Char", "Heading6Char",
            "ListParagraph", "ListBullet", "ListBullet2", "ListBullet3",
            "ListNumber", "ListNumber2", "ListNumber3", "ListNumber4", "ListNumber5",
            "NoSpacing", "Quote", "Caption", "TableCaption", "ImageCaption",
        }
        skip_ids = {"VerbatimChar", "SourceCode"}

        for style in root.findall(".//{"+ns+"}style"):
            style_id = style.get("{"+ns+"}styleId") or ""
            if style_id in skip_ids:
                continue
            name = style.find("{"+ns+"}name")
            name_val = name.get("{"+ns+"}val") if name is not None else ""
            if style_id in target_ids or name_val in target_ids or name_val.lower().startswith("heading"):
                rPr = style.find("{"+ns+"}rPr")
                if rPr is None:
                    rPr = ET.SubElement(style, "{"+ns+"}rPr")
                rFonts = rPr.find("{"+ns+"}rFonts")
                if rFonts is None:
                    rFonts = ET.SubElement(rPr, "{"+ns+"}rFonts")
                set_fonts(rFonts)

        new_data = ET.tostring(root, encoding="utf-8", xml_declaration=True)
        tmp_path = docx_path + ".tmp"
        with zipfile.ZipFile(docx_path, "r") as zin, zipfile.ZipFile(tmp_path, "w", compression=zipfile.ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                if item.filename == "word/styles.xml":
                    zout.writestr(item, new_data)
                else:
                    zout.writestr(item, zin.read(item.filename))
        os.replace(tmp_path, docx_path)
    except Exception:
        return


def patch_heading_styles(docx_path: str, enable_outline: bool, color_hex: str | None = None):
    if not docx_path or not os.path.isfile(docx_path):
        return
    ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    ET.register_namespace("w", ns)
    try:
        with zipfile.ZipFile(docx_path, "r") as z:
            if "word/styles.xml" not in z.namelist():
                return
            data = z.read("word/styles.xml")
        root = ET.fromstring(data)
        heading_ids = {f"Heading{i}" for i in range(1, 10)}
        heading_char_ids = {f"Heading{i}Char" for i in range(1, 10)}

        for style in root.findall(".//{"+ns+"}style"):
            style_id = style.get("{"+ns+"}styleId") or ""
            if style_id in heading_ids:
                # set heading color to white
                rPr = style.find("{"+ns+"}rPr")
                if rPr is None:
                    rPr = ET.SubElement(style, "{"+ns+"}rPr")
                color = rPr.find("{"+ns+"}color")
                if color_hex:
                    if color is None:
                        color = ET.SubElement(rPr, "{"+ns+"}color")
                    color.set("{"+ns+"}val", color_hex)
                    for attr in ("themeColor", "themeShade", "themeTint"):
                        key = "{"+ns+"}"+attr
                        if key in color.attrib:
                            del color.attrib[key]
                else:
                    if color is not None:
                        rPr.remove(color)
                # toggle outline level (controls collapsible headings)
                pPr = style.find("{"+ns+"}pPr")
                if pPr is None:
                    pPr = ET.SubElement(style, "{"+ns+"}pPr")
                outline = pPr.find("{"+ns+"}outlineLvl")
                if enable_outline:
                    if outline is None:
                        outline = ET.SubElement(pPr, "{"+ns+"}outlineLvl")
                    if "Heading" in style_id:
                        level = str(int(style_id.replace("Heading", "")) - 1)
                        outline.set("{"+ns+"}val", level)
                else:
                    if outline is not None:
                        pPr.remove(outline)
            if style_id in heading_char_ids:
                rPr = style.find("{"+ns+"}rPr")
                if rPr is None:
                    rPr = ET.SubElement(style, "{"+ns+"}rPr")
                color = rPr.find("{"+ns+"}color")
                if color_hex:
                    if color is None:
                        color = ET.SubElement(rPr, "{"+ns+"}color")
                    color.set("{"+ns+"}val", color_hex)
                    for attr in ("themeColor", "themeShade", "themeTint"):
                        key = "{"+ns+"}"+attr
                        if key in color.attrib:
                            del color.attrib[key]
                else:
                    if color is not None:
                        rPr.remove(color)

        new_data = ET.tostring(root, encoding="utf-8", xml_declaration=True)
        tmp_path = docx_path + ".tmp"
        with zipfile.ZipFile(docx_path, "r") as zin, zipfile.ZipFile(tmp_path, "w", compression=zipfile.ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                if item.filename == "word/styles.xml":
                    zout.writestr(item, new_data)
                else:
                    zout.writestr(item, zin.read(item.filename))
        os.replace(tmp_path, docx_path)
    except Exception:
        return


def patch_document_outline(docx_path: str, enable_outline: bool):
    if enable_outline:
        return
    if not docx_path or not os.path.isfile(docx_path):
        return
    ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    ET.register_namespace("w", ns)
    try:
        with zipfile.ZipFile(docx_path, "r") as z:
            if "word/document.xml" not in z.namelist():
                return
            data = z.read("word/document.xml")
        root = ET.fromstring(data)
        for ppr in root.findall(".//{"+ns+"}pPr"):
            outline = ppr.find("{"+ns+"}outlineLvl")
            if outline is not None:
                ppr.remove(outline)
        new_data = ET.tostring(root, encoding="utf-8", xml_declaration=True)
        tmp_path = docx_path + ".tmp"
        with zipfile.ZipFile(docx_path, "r") as zin, zipfile.ZipFile(tmp_path, "w", compression=zipfile.ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                if item.filename == "word/document.xml":
                    zout.writestr(item, new_data)
                else:
                    zout.writestr(item, zin.read(item.filename))
        os.replace(tmp_path, docx_path)
    except Exception:
        return


def demote_heading_paragraphs(docx_path: str, size_map: dict | None = None, bold_levels: set | None = None):
    if not docx_path or not os.path.isfile(docx_path):
        return
    ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    ET.register_namespace("w", ns)
    try:
        with zipfile.ZipFile(docx_path, "r") as z:
            if "word/document.xml" not in z.namelist():
                return
            data = z.read("word/document.xml")
            styles_data = z.read("word/styles.xml") if "word/styles.xml" in z.namelist() else None
        root = ET.fromstring(data)
        heading_vals = {f"Heading{i}" for i in range(1, 10)}
        if size_map is None:
            size_map = {1: 20, 2: 18, 3: 16, 4: 14, 5: 13, 6: 12}
            if styles_data:
                try:
                    styles_root = ET.fromstring(styles_data)
                    for style in styles_root.findall(".//{" + ns + "}style"):
                        style_id = style.get("{" + ns + "}styleId")
                        if not style_id or not style_id.startswith("Heading"):
                            continue
                        try:
                            level = int(style_id.replace("Heading", ""))
                        except ValueError:
                            continue
                        rpr = style.find("{" + ns + "}rPr")
                        if rpr is None:
                            continue
                        sz = rpr.find("{" + ns + "}sz")
                        if sz is None:
                            sz = rpr.find("{" + ns + "}szCs")
                        if sz is None:
                            continue
                        val = sz.get("{" + ns + "}val")
                        if val and val.isdigit():
                            size_map[level] = max(1, int(int(val) / 2))
                except Exception:
                    pass
        if bold_levels is None:
            bold_levels = {1, 2, 3, 4, 5, 6}
        for para in root.findall(".//{" + ns + "}p"):
            ppr = para.find("{" + ns + "}pPr")
            if ppr is None:
                continue
            pstyle = ppr.find("{" + ns + "}pStyle")
            if pstyle is None:
                continue
            style_val = pstyle.get("{" + ns + "}val")
            if style_val not in heading_vals:
                continue
            try:
                level = int(style_val.replace("Heading", ""))
            except ValueError:
                level = 1
            pstyle.set("{" + ns + "}val", "Normal")
            outline = ppr.find("{" + ns + "}outlineLvl")
            if outline is not None:
                ppr.remove(outline)
            size = size_map.get(level)
            if not size:
                continue
            is_bold = level in bold_levels
            for run in para.findall(".//{" + ns + "}r"):
                rpr = run.find("{" + ns + "}rPr")
                if rpr is None:
                    rpr = ET.SubElement(run, "{" + ns + "}rPr")
                sz = rpr.find("{" + ns + "}sz")
                if sz is None:
                    sz = ET.SubElement(rpr, "{" + ns + "}sz")
                sz.set("{" + ns + "}val", str(int(size * 2)))
                szcs = rpr.find("{" + ns + "}szCs")
                if szcs is None:
                    szcs = ET.SubElement(rpr, "{" + ns + "}szCs")
                szcs.set("{" + ns + "}val", str(int(size * 2)))
                bold = rpr.find("{" + ns + "}b")
                if bold is None:
                    bold = ET.SubElement(rpr, "{" + ns + "}b")
                bold.set("{" + ns + "}val", "1" if is_bold else "0")
                boldcs = rpr.find("{" + ns + "}bCs")
                if boldcs is None:
                    boldcs = ET.SubElement(rpr, "{" + ns + "}bCs")
                boldcs.set("{" + ns + "}val", "1" if is_bold else "0")
        new_data = ET.tostring(root, encoding="utf-8", xml_declaration=True)
        tmp_path = docx_path + ".tmp"
        with zipfile.ZipFile(docx_path, "r") as zin, zipfile.ZipFile(tmp_path, "w", compression=zipfile.ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                if item.filename == "word/document.xml":
                    zout.writestr(item, new_data)
                else:
                    zout.writestr(item, zin.read(item.filename))
        os.replace(tmp_path, docx_path)
    except Exception:
        return
def remove_literal_math_markers(docx_path: str):
    if not docx_path or not os.path.isfile(docx_path):
        return
    markers = {"$$", r"\[", r"\]"}
    try:
        doc = Document(docx_path)
        for paragraph in list(doc.paragraphs):
            if paragraph.text.strip() in markers:
                paragraph._element.getparent().remove(paragraph._element)
        doc.save(docx_path)
    except Exception:
        return


def table_to_plaintext(text: str) -> str:
    lines = text.splitlines()
    out = []
    i = 0
    while i < len(lines):
        line = lines[i]
        if line.count("|") >= 2:
            # collect table block
            block = []
            while i < len(lines) and lines[i].count("|") >= 2:
                block.append(lines[i])
                i += 1
            # drop separator lines
            rows = []
            for bl in block:
                stripped = bl.strip()
                if re.fullmatch(r"\|?[\s:-]+\|?[\s:-]*", stripped):
                    continue
                parts = [p.strip() for p in stripped.strip("|").split("|")]
                rows.append(parts)
            # compute column widths
            col_count = max((len(r) for r in rows), default=0)
            widths = [0] * col_count
            for r in rows:
                for idx in range(col_count):
                    cell = r[idx] if idx < len(r) else ""
                    widths[idx] = max(widths[idx], len(cell))
            # emit plain text table
            for r in rows:
                padded = []
                for idx in range(col_count):
                    cell = r[idx] if idx < len(r) else ""
                    padded.append(cell.ljust(widths[idx]))
                out.append("  ".join(padded).rstrip())
            continue
        out.append(line)
        i += 1
    return "\n".join(out)


class DropArea(QtWidgets.QFrame):
    fileDropped = QtCore.pyqtSignal(list)
    filePicked = QtCore.pyqtSignal()

    def __init__(self, parent=None):
        super().__init__(parent)
        self.setAcceptDrops(True)
        self.setObjectName("DropArea")
        self.setCursor(QtGui.QCursor(QtCore.Qt.CursorShape.PointingHandCursor))
        label_main = QtWidgets.QLabel("拖拽文件到这里")
        label_main.setAlignment(QtCore.Qt.AlignmentFlag.AlignCenter)
        label_main.setObjectName("DropLabelMain")
        label_sub = QtWidgets.QLabel("支持 .md / .txt")
        label_sub.setAlignment(QtCore.Qt.AlignmentFlag.AlignCenter)
        label_sub.setObjectName("DropLabelSub")
        layout = QtWidgets.QVBoxLayout(self)
        layout.addStretch(1)
        layout.addWidget(label_main)
        layout.addWidget(label_sub)
        layout.addStretch(1)

    def mouseDoubleClickEvent(self, event):
        if event.button() == QtCore.Qt.MouseButton.LeftButton:
            self.filePicked.emit()
        super().mouseDoubleClickEvent(event)

    def dragEnterEvent(self, event):
        if event.mimeData().hasUrls():
            self.setProperty("dragging", True)
            self.style().unpolish(self)
            self.style().polish(self)
            event.acceptProposedAction()

    def dragLeaveEvent(self, event):
        self.setProperty("dragging", False)
        self.style().unpolish(self)
        self.style().polish(self)
        super().dragLeaveEvent(event)

    def dropEvent(self, event):
        urls = event.mimeData().urls()
        paths = [u.toLocalFile() for u in urls if u.isLocalFile()]
        if paths:
            self.fileDropped.emit(paths)
        self.setProperty("dragging", False)
        self.style().unpolish(self)
        self.style().polish(self)
        event.acceptProposedAction()

    def enterEvent(self, event):
        self.setProperty("hover", True)
        self.style().unpolish(self)
        self.style().polish(self)
        super().enterEvent(event)

    def leaveEvent(self, event):
        self.setProperty("hover", False)
        self.style().unpolish(self)
        self.style().polish(self)
        super().leaveEvent(event)


class ExportList(QtWidgets.QListWidget):
    itemRemoved = QtCore.pyqtSignal(str)
    def __init__(self, parent=None):
        super().__init__(parent)
        self.setObjectName("ExportList")
        self.setSpacing(12)
        self.setSelectionMode(QtWidgets.QAbstractItemView.SelectionMode.NoSelection)
        self.setVerticalScrollMode(QtWidgets.QAbstractItemView.ScrollMode.ScrollPerPixel)
        self.setAcceptDrops(True)
        self.setDragDropMode(QtWidgets.QAbstractItemView.DragDropMode.DropOnly)

    def add_item(self, path: str):
        clean_path = os.path.normpath(str(path).strip())
        key = os.path.normcase(clean_path)
        item = QtWidgets.QListWidgetItem()
        item.setData(QtCore.Qt.ItemDataRole.UserRole, {"path": clean_path, "key": key})
        widget = self._build_item_widget(clean_path)
        item.setSizeHint(widget.sizeHint())
        self.addItem(item)
        self.setItemWidget(item, widget)

    def _build_item_widget(self, path: str):
        box = QtWidgets.QFrame()
        box.setObjectName("ExportItem")
        layout = QtWidgets.QHBoxLayout(box)
        layout.setContentsMargins(10, 9, 10, 9)
        layout.setSpacing(8)

        name = QtWidgets.QLabel(os.path.basename(path))
        name.setObjectName("ExportName")
        name.setToolTip(path)
        name.setWordWrap(False)
        name.setMinimumWidth(0)
        name.setSizePolicy(QtWidgets.QSizePolicy.Policy.Expanding, QtWidgets.QSizePolicy.Policy.Preferred)
        layout.addWidget(name, 1)

        layout.addWidget(self._mini_btn("📄", lambda: self._open_file(path)))
        layout.addWidget(self._mini_btn("📁", lambda: self._open_folder(path)))
        layout.addWidget(self._mini_btn("🗑", lambda: self._delete_path(path)))
        return box

    def _mini_btn(self, text, handler):
        btn = QtWidgets.QPushButton(text)
        btn.setObjectName("MiniBtn")
        btn.setFixedSize(30, 26)
        if any(ord(ch) > 127 for ch in text):
            btn.setFont(QtGui.QFont("Segoe UI Emoji", 9))
        def safe_handler():
            try:
                handler()
            except Exception:
                QtWidgets.QMessageBox.warning(self, "提示", "操作失败。")
        btn.clicked.connect(safe_handler)
        return btn

    def _delete_path(self, path: str):
        reply = QtWidgets.QMessageBox.question(self, "确认删除", "确定要删除这个文件吗？")
        if reply != QtWidgets.QMessageBox.StandardButton.Yes:
            return
        remove_error = None
        try:
            clean_path = os.path.normpath(path)
            if os.path.exists(clean_path):
                os.remove(clean_path)
        except Exception as exc:
            remove_error = exc
        self._remove_all_items(path)
        if remove_error is not None:
            QtWidgets.QMessageBox.warning(self, "提示", "删除失败。")

    def _remove_all_items(self, path: str):
        key = os.path.normcase(os.path.normpath(path))
        removed = False
        for row in range(self.count() - 1, -1, -1):
            item = self.item(row)
            payload = item.data(QtCore.Qt.ItemDataRole.UserRole) or {}
            if payload.get("key") == key:
                self.takeItem(row)
                removed = True
        if removed:
            self.itemRemoved.emit(path)

    def _open_file(self, path: str):
        clean_path = os.path.normpath(path)
        if not os.path.exists(clean_path):
            QtWidgets.QMessageBox.warning(self, "提示", "文件不存在。")
            return
        os.startfile(clean_path)

    def _open_folder(self, path: str):
        clean_path = os.path.normpath(path)
        if os.path.exists(clean_path):
            subprocess.run(["explorer", "/select,", clean_path], check=False)
            return
        folder = os.path.dirname(clean_path)
        if os.path.isdir(folder):
            os.startfile(folder)

    def dragEnterEvent(self, event):
        if event.mimeData().hasFormat("application/x-font-preset"):
            event.acceptProposedAction()
            return
        super().dragEnterEvent(event)

    def dropEvent(self, event):
        if not event.mimeData().hasFormat("application/x-font-preset"):
            super().dropEvent(event)
            return
        data = bytes(event.mimeData().data("application/x-font-preset")).decode("utf-8", errors="ignore")
        try:
            payload = json.loads(data)
        except Exception:
            return
        pos = event.position().toPoint()
        item = self.itemAt(pos)
        if not item:
            return
        payload_item = item.data(QtCore.Qt.ItemDataRole.UserRole) or {}
        path = payload_item.get("path")
        if not path:
            return
        self._apply_preset_to_path(path, payload)
        event.acceptProposedAction()

    def _apply_preset_to_path(self, path: str, payload: dict):
        try:
            east = payload.get("east_asia") or "宋体"
            latin = payload.get("latin") or "Times New Roman"
            patch_styles_fonts(path, east, latin)
            patch_numbering_fonts(path, east, latin)
        except Exception:
            QtWidgets.QMessageBox.warning(self, "提示", "应用字体失败。")


class PresetList(QtWidgets.QListWidget):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.setObjectName("PresetList")
        self.setDragEnabled(True)
        self.setDragDropMode(QtWidgets.QAbstractItemView.DragDropMode.DragOnly)
        self.setSelectionMode(QtWidgets.QAbstractItemView.SelectionMode.SingleSelection)
        self.setVerticalScrollMode(QtWidgets.QAbstractItemView.ScrollMode.ScrollPerPixel)

    def startDrag(self, supportedActions):
        item = self.currentItem()
        if not item:
            return
        payload = item.data(QtCore.Qt.ItemDataRole.UserRole)
        if not payload:
            return
        mime = QtCore.QMimeData()
        mime.setData("application/x-font-preset", json.dumps(payload, ensure_ascii=False).encode("utf-8"))
        drag = QtGui.QDrag(self)
        drag.setMimeData(mime)
        drag.exec(supportedActions)


class GlowOverlay(QtWidgets.QWidget):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.setAttribute(QtCore.Qt.WidgetAttribute.WA_TransparentForMouseEvents, True)
        self.setAttribute(QtCore.Qt.WidgetAttribute.WA_NoSystemBackground, True)
        self.setAttribute(QtCore.Qt.WidgetAttribute.WA_TranslucentBackground, True)
        self._pos = None

    def set_pos(self, pos):
        self._pos = pos
        self.update()

    def paintEvent(self, event):
        if not self._pos:
            return
        painter = QtGui.QPainter(self)
        painter.setRenderHint(QtGui.QPainter.RenderHint.Antialiasing)
        gradient = QtGui.QRadialGradient(self._pos, 170)
        gradient.setColorAt(0.0, QtGui.QColor(120, 220, 255, 60))
        gradient.setColorAt(1.0, QtGui.QColor(0, 0, 0, 0))
        painter.fillRect(self.rect(), gradient)


class WorkerSignals(QtCore.QObject):
    finished = QtCore.pyqtSignal(object)
    error = QtCore.pyqtSignal(str, str)


class BackgroundTask(QtCore.QRunnable):
    def __init__(self, func):
        super().__init__()
        self.func = func
        self.signals = WorkerSignals()

    def run(self):
        try:
            result = self.func()
            self.signals.finished.emit(result)
        except Exception as exc:
            self.signals.error.emit("处理失败", str(exc))


class TitleBar(QtWidgets.QFrame):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.setObjectName("TitleBar")
        self.setFixedHeight(42)
        self._window = parent.window() if parent is not None else None
        layout = QtWidgets.QHBoxLayout(self)
        layout.setContentsMargins(14, 6, 10, 6)
        layout.setSpacing(10)

        self.icon_label = QtWidgets.QLabel()
        self.icon_label.setFixedSize(18, 18)
        self.icon_label.setScaledContents(True)

        title = ""
        if self._window is not None:
            title = self._window.windowTitle()
        self.title_label = QtWidgets.QLabel(title)
        self.title_label.setObjectName("TitleLabel")

        layout.addWidget(self.icon_label)
        layout.addWidget(self.title_label)
        layout.addStretch(1)

        self.sub_label = QtWidgets.QLabel("Markdown 转 Word")
        self.sub_label.setObjectName("SubTitle")
        layout.addWidget(self.sub_label)

        self.btn_min = QtWidgets.QToolButton()
        self.btn_min.setObjectName("TitleMin")
        self.btn_min.setText("—")
        self.btn_min.setToolTip("最小化")

        self.btn_max = QtWidgets.QToolButton()
        self.btn_max.setObjectName("TitleMax")
        self.btn_max.setText("□")
        self.btn_max.setToolTip("最大化/还原")

        self.btn_close = QtWidgets.QToolButton()
        self.btn_close.setObjectName("TitleClose")
        self.btn_close.setText("×")
        self.btn_close.setToolTip("关闭")

        layout.addWidget(self.btn_min)
        layout.addWidget(self.btn_max)
        layout.addWidget(self.btn_close)

        if self._window is not None:
            self.btn_min.clicked.connect(self._window.showMinimized)
            if hasattr(self._window, "toggle_maximize"):
                self.btn_max.clicked.connect(self._window.toggle_maximize)
            self.btn_close.clicked.connect(self._window.close)

    def set_icon(self, icon: QtGui.QIcon):
        if icon and not icon.isNull():
            self.icon_label.setPixmap(icon.pixmap(18, 18))

    def update_max_icon(self, maximized: bool):
        self.btn_max.setText("❐" if maximized else "□")

    def is_in_drag_area(self, pos: QtCore.QPoint) -> bool:
        for btn in (self.btn_min, self.btn_max, self.btn_close):
            if btn.geometry().contains(pos):
                return False
        return True

    def mousePressEvent(self, event):
        if event.button() == QtCore.Qt.MouseButton.LeftButton:
            pos = event.position().toPoint()
            if self.is_in_drag_area(pos):
                window = self.window().windowHandle()
                if window is not None:
                    window.startSystemMove()
                event.accept()
                return
        super().mousePressEvent(event)

    def mouseDoubleClickEvent(self, event):
        if event.button() == QtCore.Qt.MouseButton.LeftButton:
            pos = event.position().toPoint()
            if self.is_in_drag_area(pos) and self._window is not None and hasattr(self._window, "toggle_maximize"):
                self._window.toggle_maximize()
        super().mouseDoubleClickEvent(event)


class MainWindow(QtWidgets.QMainWindow):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("团子转换器")
        self.resize(1100, 680)
        self.pandoc = find_pandoc()
        self.glow = None
        self.data_dir = get_user_data_dir()
        if getattr(sys, "frozen", False):
            base_dir = os.path.dirname(sys.executable)
            self.log_dir = os.path.join(base_dir, "logs")
        else:
            self.log_dir = os.path.join(self.data_dir, "logs")
        os.makedirs(self.log_dir, exist_ok=True)
        stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        self.log_file = os.path.join(self.log_dir, f"run_{stamp}.log")
        self.presets_file = os.path.join(os.path.dirname(__file__), "templates", "presets.json")
        self.history_file = os.path.join(self.data_dir, "export_history.json")
        self.settings_file = os.path.join(self.data_dir, "settings.json")
        self.app_icon = None
        icon_path = os.path.join(os.path.dirname(__file__), "assets", "app.ico")
        if os.path.isfile(icon_path):
            self.app_icon = QtGui.QIcon(icon_path)
            self.setWindowIcon(self.app_icon)

        self.source_mode = "file"
        self.open_after_export = False
        self.keep_rules = True
        self.enable_log = False
        self.use_custom_fonts = False
        self.enable_heading_outline = False
        self.custom_name_enabled = False
        self.enable_page_numbers = True
        self.page_start = 1
        self.page_start_page = 1
        self.output_dir = ""
        self._export_history = []
        self.thread_pool = QtCore.QThreadPool.globalInstance()
        self._busy_count = 0
        self.font_presets = [
            ("宋体 + Times New Roman（学术/正式）", "academic_songti_times.docx"),
            ("微软雅黑 + Calibri（通用/清晰）", "modern_yahei_calibri.docx"),
            ("仿宋 + Times New Roman（公文/报告）", "official_fangsong_times.docx"),
            ("黑体 + Arial（标题醒目）", "bold_heiti_arial.docx"),
            ("楷体 + Times New Roman（演讲/文案）", "calligraphy_kaiti_times.docx"),
        ]
        self.bold_presets = [
            ("学术/论文（仅 1-2 级加粗）", {1, 2}),
            ("报告/公文（1-3 级加粗）", {1, 2, 3}),
            ("教程/说明（1-4 级加粗）", {1, 2, 3, 4}),
            ("全部加粗（1-6 级）", {1, 2, 3, 4, 5, 6}),
            ("不加粗（仅字号）", set()),
        ]
        self.bold_levels = {1, 2}
        self.font_name_map = {
            "academic_songti_times.docx": ("宋体", "Times New Roman"),
            "modern_yahei_calibri.docx": ("微软雅黑", "Calibri"),
            "official_fangsong_times.docx": ("仿宋", "Times New Roman"),
            "bold_heiti_arial.docx": ("黑体", "Arial"),
            "calligraphy_kaiti_times.docx": ("楷体", "Times New Roman"),
            "reference.docx": ("宋体", "Times New Roman"),
        }
        self.font_zh_list = [
            ("宋体（学术/正式）", "宋体"),
            ("微软雅黑（通用/清晰）", "微软雅黑"),
            ("仿宋（公文/报告）", "仿宋"),
            ("黑体（标题醒目）", "黑体"),
            ("楷体（演讲/文案）", "楷体"),
            ("等线（现代/简洁）", "等线"),
            ("华文中宋（论文/正文）", "华文中宋"),
            ("华文宋体（书面）", "华文宋体"),
            ("华文仿宋（公文）", "华文仿宋"),
            ("华文楷体（文案）", "华文楷体"),
            ("思源黑体（无衬线）", "思源黑体"),
            ("思源宋体（衬线）", "思源宋体"),
            ("霞鹜文楷（文艺）", "霞鹜文楷"),
        ]
        self.font_en_list = [
            ("Times New Roman（学术/正式）", "Times New Roman"),
            ("Calibri（通用/清晰）", "Calibri"),
            ("Arial（标题醒目）", "Arial"),
            ("Cambria（阅读/书面）", "Cambria"),
            ("Georgia（文章/书籍）", "Georgia"),
            ("Garamond（经典/书籍）", "Garamond"),
            ("Palatino Linotype（文艺/正文）", "Palatino Linotype"),
            ("Segoe UI（现代/系统）", "Segoe UI"),
            ("Verdana（屏幕/清晰）", "Verdana"),
            ("Tahoma（紧凑/清晰）", "Tahoma"),
            ("Century Schoolbook（论文）", "Century Schoolbook"),
        ]
        self.ai_last_digest = None
        self.ai_last_blocks = []
        self.ai_task_text = None
        self.ai_task_blocks = []
        self.ai_task_digest = None
        self.ai_task_mode = None
        self.ai_task_output_dir = None
        self.ai_task_base_override = None
        self.ai_enabled = False
        self.work_mode = "convert"

        sys.excepthook = self._handle_exception
        self._build_ui()

    def _build_ui(self):
        central = QtWidgets.QWidget()
        central.setObjectName("Central")
        self.setCentralWidget(central)

        main = QtWidgets.QHBoxLayout(central)
        main.setContentsMargins(18, 18, 18, 18)
        main.setSpacing(18)

        left = self._build_left()
        center = self._build_center()
        right = self._build_right()

        main.addWidget(left, 3)
        main.addWidget(center, 2)
        main.addWidget(right, 3)

        self._apply_styles()
        self._apply_shadows([left, center, right])
        self._load_presets()
        self._toggle_custom_name(QtCore.Qt.CheckState.Unchecked)
        self._load_export_history()
        self._load_settings()

    def _build_left(self):
        box = QtWidgets.QFrame()
        box.setObjectName("Panel")
        layout = QtWidgets.QVBoxLayout(box)
        layout.setContentsMargins(18, 18, 18, 18)
        layout.setSpacing(14)

        header = QtWidgets.QHBoxLayout()
        header.setContentsMargins(0, 0, 0, 0)
        title = QtWidgets.QLabel("导入区")
        title.setObjectName("PanelTitle")
        header.addWidget(title)
        header.addStretch(1)

        seg = QtWidgets.QFrame()
        seg.setObjectName("Segment")
        seg_layout = QtWidgets.QHBoxLayout(seg)
        seg_layout.setContentsMargins(3, 3, 3, 3)
        seg_layout.setSpacing(6)
        self.btn_file = QtWidgets.QPushButton("文档导入")
        self.btn_text = QtWidgets.QPushButton("文本导入")
        for btn in (self.btn_file, self.btn_text):
            btn.setCheckable(True)
            btn.setObjectName("SegmentBtn")
        self.btn_file.setChecked(True)
        self.btn_file.clicked.connect(lambda: self._set_source_mode("file"))
        self.btn_text.clicked.connect(lambda: self._set_source_mode("text"))
        seg_layout.addWidget(self.btn_file)
        seg_layout.addWidget(self.btn_text)

        header.addWidget(seg)
        layout.addLayout(header)

        self.drop_area = DropArea()
        self.drop_area.fileDropped.connect(self._handle_drop)
        self.drop_area.filePicked.connect(self._pick_files)

        self.text_edit = QtWidgets.QPlainTextEdit()
        self.text_edit.setPlaceholderText("粘贴 Markdown 内容")
        self.text_edit.setObjectName("TextInput")
        self.text_edit.textChanged.connect(self._update_text_preview)

        self.stack = QtWidgets.QStackedLayout()
        self.stack.addWidget(self.drop_area)
        text_wrap = QtWidgets.QFrame()
        text_wrap.setObjectName("InnerCard")
        text_layout = QtWidgets.QVBoxLayout(text_wrap)
        text_layout.setContentsMargins(10, 10, 10, 10)
        text_layout.addWidget(self.text_edit)
        self.stack.addWidget(text_wrap)
        layout.addLayout(self.stack)
        self.btn_export_text = QtWidgets.QPushButton("导出 Word")
        self.btn_export_text.setObjectName("PrimaryBtn")
        self.btn_export_text.setMinimumHeight(36)
        icon_path = os.path.join(os.path.dirname(__file__), "assets", "export_word.svg")
        if os.path.isfile(icon_path):
            self.btn_export_text.setIcon(QtGui.QIcon(icon_path))
            self.btn_export_text.setIconSize(QtCore.QSize(16, 16))
        self.btn_export_text.clicked.connect(self._export_from_text)
        self.btn_export_text.setVisible(False)
        btn_row = QtWidgets.QHBoxLayout()
        btn_row.addStretch(1)
        btn_row.addWidget(self.btn_export_text)
        btn_row.addStretch(1)
        layout.addLayout(btn_row)
        return box

    def _build_center(self):
        box = QtWidgets.QFrame()
        box.setObjectName("Panel")
        layout = QtWidgets.QVBoxLayout(box)
        layout.setContentsMargins(18, 18, 18, 18)
        layout.setSpacing(14)

        title = QtWidgets.QLabel("可选功能")
        title.setObjectName("PanelTitle")
        layout.addWidget(title)

        self.chk_rules = QtWidgets.QCheckBox("保留分隔线")
        self.chk_fold = QtWidgets.QCheckBox("标题可折叠")
        self.chk_custom_name = QtWidgets.QCheckBox("自定义导出文件名")
        self.chk_rules.setObjectName("RuleCheck")
        self.chk_fold.setObjectName("FoldCheck")
        self.chk_custom_name.setObjectName("NameCheck")
        self.chk_page = QtWidgets.QCheckBox("自动页码")
        self.chk_page.setObjectName("PageCheck")
        self.chk_rules.setChecked(True)
        self.chk_fold.setChecked(False)
        self.chk_custom_name.setChecked(False)
        self.chk_page.setChecked(True)
        self.chk_rules.stateChanged.connect(lambda v: setattr(self, "keep_rules", v == QtCore.Qt.CheckState.Checked))
        self.chk_fold.stateChanged.connect(lambda v: setattr(self, "enable_heading_outline", v == QtCore.Qt.CheckState.Checked))
        self.chk_custom_name.toggled.connect(self._toggle_custom_name)
        self.chk_page.toggled.connect(self._toggle_page_numbers)

        layout.addWidget(self.chk_rules)
        layout.addWidget(self.chk_fold)
        layout.addWidget(self.chk_custom_name)
        layout.addWidget(self.chk_page)

        self.name_edit = QtWidgets.QLineEdit()
        self.name_edit.setObjectName("TextInput")
        self.name_edit.setPlaceholderText("原文件名已转换")
        self.name_edit.setClearButtonEnabled(True)
        layout.addWidget(self.name_edit)

        layout.addSpacing(6)
        dir_label = QtWidgets.QLabel("导出文件夹")
        dir_label.setObjectName("LogTitle")
        layout.addWidget(dir_label)
        dir_row = QtWidgets.QHBoxLayout()
        dir_row.setContentsMargins(0, 0, 0, 0)
        self.output_dir_edit = QtWidgets.QLineEdit()
        self.output_dir_edit.setObjectName("TextInput")
        self.output_dir_edit.setPlaceholderText("默认桌面")
        self.output_dir_edit.setReadOnly(False)
        self.output_dir_edit.setEnabled(True)
        self.output_dir_edit.editingFinished.connect(self._commit_output_dir_text)
        self.btn_output_dir = QtWidgets.QPushButton("选择")
        self.btn_output_dir.setObjectName("SecondaryBtn")
        self.btn_output_dir.setEnabled(True)
        self.btn_output_dir.clicked.connect(self._pick_output_dir)
        dir_row.addWidget(self.output_dir_edit, 1)
        dir_row.addWidget(self.btn_output_dir)
        layout.addLayout(dir_row)

        show_row = QtWidgets.QHBoxLayout()
        show_row.setContentsMargins(0, 0, 0, 0)
        show_label = QtWidgets.QLabel("页码显示从")
        show_label.setObjectName("LogTitle")
        self.page_show_spin = QtWidgets.QSpinBox()
        self.page_show_spin.setObjectName("TextInput")
        self.page_show_spin.setRange(1, 2)
        self.page_show_spin.setValue(1)
        self.page_show_spin.setToolTip("从第几页开始显示页码（支持第1或第2页）")
        self.page_show_spin.valueChanged.connect(lambda v: setattr(self, "page_start_page", int(v)))
        show_row.addWidget(show_label)
        show_row.addStretch(1)
        show_row.addWidget(self.page_show_spin)
        layout.addLayout(show_row)

        page_row = QtWidgets.QHBoxLayout()
        page_row.setContentsMargins(0, 0, 0, 0)
        page_label = QtWidgets.QLabel("页码起始")
        page_label.setObjectName("LogTitle")
        self.page_spin = QtWidgets.QSpinBox()
        self.page_spin.setObjectName("TextInput")
        self.page_spin.setRange(1, 999)
        self.page_spin.setValue(1)
        self.page_spin.setToolTip("页码从该数值开始，例如 2 表示第一页显示为 2")
        self.page_spin.valueChanged.connect(lambda v: setattr(self, "page_start", int(v)))
        page_row.addWidget(page_label)
        page_row.addStretch(1)
        page_row.addWidget(self.page_spin)
        layout.addLayout(page_row)
        self._toggle_page_numbers(True)

        font_label = QtWidgets.QLabel("字体选择")
        font_label.setObjectName("LogTitle")
        self.font_combo = QtWidgets.QComboBox()
        self.font_combo.setObjectName("TextInput")
        for label, filename in self.font_presets:
            self.font_combo.addItem(label, filename)
        layout.addWidget(font_label)
        layout.addWidget(self.font_combo)

        bold_label = QtWidgets.QLabel("加粗预设")
        bold_label.setObjectName("LogTitle")
        self.bold_combo = QtWidgets.QComboBox()
        self.bold_combo.setObjectName("TextInput")
        for label, levels in self.bold_presets:
            self.bold_combo.addItem(label, levels)
        self.bold_combo.currentIndexChanged.connect(self._update_bold_preset)
        layout.addWidget(bold_label)
        layout.addWidget(self.bold_combo)
        layout.addStretch(1)

        return box

    def _update_bold_preset(self, idx: int):
        data = self.bold_combo.currentData() if hasattr(self, "bold_combo") else None
        if isinstance(data, set):
            self.bold_levels = data
        elif isinstance(data, (list, tuple)):
            self.bold_levels = set(data)
        else:
            self.bold_levels = {1, 2}

    def _build_right(self):
        box = QtWidgets.QFrame()
        box.setObjectName("Panel")
        layout = QtWidgets.QVBoxLayout(box)
        layout.setContentsMargins(18, 18, 18, 18)
        layout.setSpacing(14)

        header = QtWidgets.QHBoxLayout()
        header.setContentsMargins(0, 0, 0, 0)
        title = QtWidgets.QLabel("导出区")
        title.setObjectName("PanelTitle")
        header.addWidget(title)
        header.addStretch(1)
        subtitle = QtWidgets.QLabel("Markdown 转 Word")
        subtitle.setObjectName("SubCorner")
        header.addWidget(subtitle)
        layout.addLayout(header)

        self.export_list = ExportList()
        self.export_list.itemRemoved.connect(self._remove_history)
        export_wrap = QtWidgets.QFrame()
        export_wrap.setObjectName("InnerCard")
        export_layout = QtWidgets.QVBoxLayout(export_wrap)
        export_layout.setContentsMargins(12, 12, 12, 12)
        export_layout.addWidget(self.export_list)

        self.output_text = QtWidgets.QPlainTextEdit()
        self.output_text.setObjectName("TextOutput")
        self.output_text.setReadOnly(True)
        output_wrap = QtWidgets.QFrame()
        output_wrap.setObjectName("InnerCard")
        output_layout = QtWidgets.QVBoxLayout(output_wrap)
        output_layout.setContentsMargins(12, 12, 12, 12)
        output_layout.addWidget(self.output_text)

        self.right_stack = QtWidgets.QStackedLayout()
        self.right_stack.addWidget(export_wrap)
        self.right_stack.addWidget(output_wrap)
        layout.addLayout(self.right_stack)
        return box

    def _apply_styles(self):
        QtWidgets.QApplication.setFont(QtGui.QFont("Microsoft YaHei UI", 10))
        self.setStyleSheet(
            """
            QWidget#Central {
                background: qlineargradient(x1:0, y1:0, x2:0, y2:1,
                stop:0 #E2EBE4, stop:0.55 #DCE5DE, stop:1 #E3D6C0);
            }
            QFrame#Panel {
                background: qlineargradient(x1:0, y1:0, x2:0, y2:1,
                stop:0 rgba(255,255,255,0.70),
                stop:0.10 rgba(255,255,255,0.48),
                stop:1 rgba(255,255,255,0.44));
                border: 1px solid rgba(184, 200, 190, 0.85);
                border-radius: 20px;
            }
            QLabel#PanelTitle { color: #3F4A52; font-size: 17px; font-weight: 600; letter-spacing: 0.3px; }
            QLabel#SubCorner { color: #7B8A95; font-size: 11px; padding-top: 2px; }
            QFrame#InnerCard {
                background: qlineargradient(x1:0, y1:0, x2:0, y2:1,
                stop:0 rgba(236, 242, 236, 0.58),
                stop:0.12 rgba(230, 236, 230, 0.48),
                stop:1 rgba(225, 231, 225, 0.44));
                border: 1px solid rgba(176, 192, 182, 0.9);
                border-radius: 16px;
            }
            QFrame#Segment {
                background: rgba(255,255,255,0.62);
                border: 1px solid rgba(192, 206, 196, 0.85);
                border-radius: 14px;
            }
            QPushButton#SegmentBtn {
                color: #4A5861; background: transparent; border: none;
                padding: 7px 16px; border-radius: 11px;
            }
            QPushButton#SegmentBtn:hover {
                background: rgba(255,255,255,0.75);
                border: 1px solid rgba(169, 199, 221, 0.7);
            }
            QPushButton#SegmentBtn:checked {
                background: qlineargradient(x1:0, y1:0, x2:1, y2:1,
                stop:0 rgba(191, 221, 242, 0.95), stop:1 rgba(169, 199, 221, 0.85));
                color: #2F3A40;
            }
            QFrame#DropArea {
                background: rgba(230, 236, 230, 0.38);
                border: 1px dashed rgba(159, 182, 200, 0.65);
                border-radius: 18px;
            }
            QFrame#DropArea[hover="true"], QFrame#DropArea[dragging="true"] {
                background: rgba(230, 236, 230, 0.50);
                border: 1px dashed #A9C7DD;
            }
            QLabel#DropLabelMain { color: #3F4A52; font-size: 14px; }
            QLabel#DropLabelSub { color: #7B8A95; font-size: 12px; }
            QPlainTextEdit#TextInput {
                background: transparent;
                color: #2F3A40;
                border: none;
                border-radius: 14px; padding: 12px;
            }
            QPlainTextEdit#TextInput::viewport {
                background: transparent;
                border-radius: 14px;
            }
            QLineEdit#TextInput {
                background: rgba(255,255,255,0.80);
                color: #2F3A40;
                border: 1px solid rgba(192, 206, 196, 0.95);
                border-radius: 12px; padding: 6px 10px;
            }
            QPlainTextEdit#TextOutput {
                background: transparent;
                color: #2F3A40;
                border: none;
                border-radius: 14px; padding: 12px;
            }
            QPlainTextEdit#TextOutput::viewport {
                background: transparent;
                border-radius: 14px;
            }
            QCheckBox { color: #3F4A52; font-size: 14px; spacing: 8px; }
            QCheckBox::indicator {
                width: 18px; height: 18px;
                border-radius: 6px;
                border: 1px solid rgba(178, 196, 188, 0.9);
                background: rgba(255,255,255,0.8);
            }
            QCheckBox::indicator:checked {
                background: qlineargradient(x1:0, y1:0, x2:1, y2:1,
                stop:0 rgba(191, 227, 178, 0.95), stop:1 rgba(169, 207, 162, 0.9));
                border: 1px solid rgba(169, 207, 162, 0.9);
            }
            QCheckBox#RuleCheck::indicator:checked {
                background: qlineargradient(x1:0, y1:0, x2:1, y2:1,
                stop:0 rgba(191, 227, 178, 0.95), stop:1 rgba(169, 207, 162, 0.9));
                border: 1px solid rgba(169, 207, 162, 0.9);
            }
            QCheckBox#FoldCheck::indicator:checked {
                background: qlineargradient(x1:0, y1:0, x2:1, y2:1,
                stop:0 rgba(191, 221, 242, 0.95), stop:1 rgba(169, 199, 221, 0.9));
                border: 1px solid rgba(169, 199, 221, 0.9);
            }
            QCheckBox#NameCheck::indicator:checked {
                background: qlineargradient(x1:0, y1:0, x2:1, y2:1,
                stop:0 rgba(243, 225, 166, 0.95), stop:1 rgba(229, 208, 146, 0.9));
                border: 1px solid rgba(229, 208, 146, 0.9);
            }
            QCheckBox#PageCheck::indicator:checked {
                background: qlineargradient(x1:0, y1:0, x2:1, y2:1,
                stop:0 rgba(215, 195, 203, 0.95), stop:1 rgba(205, 180, 191, 0.9));
                border: 1px solid rgba(202, 176, 187, 0.9);
            }
            QListWidget#ExportList {
                background: rgba(225, 232, 225, 0.35);
                border: 1px solid rgba(186, 200, 190, 0.9);
                border-radius: 12px; color: #2F3A40;
            }
            QFrame#ExportItem {
                background: rgba(255,255,255,0.72);
                border: 1px solid rgba(192, 206, 196, 0.95);
                border-radius: 14px;
            }
            QFrame#ExportItem:hover {
                border: 1px solid rgba(169, 199, 221, 0.75);
                background: rgba(255,255,255,0.82);
            }
            QLabel#ExportName {
                color: #2F3A40;
            }
            QPushButton#MiniBtn {
                background: rgba(255,255,255,0.72);
                color: #3F4A52;
                border: 1px solid rgba(192, 206, 196, 0.95);
                border-radius: 12px; padding: 0px;
                min-width: 28px; min-height: 24px;
                font-size: 12px;
            }
            QPushButton#MiniBtn:hover {
                background: rgba(255,255,255,0.86);
                border: 1px solid rgba(229, 208, 146, 0.9);
            }
            QPushButton#PrimaryBtn {
                background: qlineargradient(x1:0, y1:0, x2:1, y2:1,
                stop:0 rgba(191, 221, 242, 0.95), stop:1 rgba(169, 207, 162, 0.9));
                color: #2F3A40;
                border: 1px solid rgba(169, 199, 221, 0.8);
                border-radius: 14px;
                padding: 6px 14px;
            }
            QPushButton#PrimaryBtn:hover {
                background: qlineargradient(x1:0, y1:0, x2:1, y2:1,
                stop:0 rgba(191, 221, 242, 1.0), stop:1 rgba(169, 207, 162, 0.95));
                border: 1px solid rgba(229, 208, 146, 0.95);
            }
            QFrame#LogBox {
                background: rgba(255,255,255,0.62);
                border: 1px solid rgba(200, 212, 202, 0.85);
                border-radius: 14px;
            }
            QLabel#LogTitle { color: #7B8A95; font-size: 12px; }
            QPlainTextEdit#LogArea {
                background: transparent; border: none; color: #7B8A95;
                font-size: 12px;
            }
            QComboBox#TextInput {
                background: rgba(255,255,255,0.80);
                color: #2F3A40;
                border: 1px solid rgba(192, 206, 196, 0.95);
                border-radius: 12px; padding: 6px 10px;
            }
            QComboBox#TextInput::drop-down {
                border: none;
            }
            QComboBox#TextInput QAbstractItemView {
                background: rgba(244, 248, 244, 0.98);
                color: #2F3A40;
                border: 1px solid rgba(192, 206, 196, 0.95);
                border-radius: 10px;
                selection-background-color: rgba(169, 199, 221, 0.55);
                selection-color: #2F3A40;
                outline: 0;
            }
            QPushButton#SecondaryBtn {
                background: rgba(255,255,255,0.70);
                color: #3F4A52;
                border: 1px solid rgba(192, 206, 196, 0.95);
                border-radius: 12px;
                padding: 4px 12px;
                min-height: 28px;
            }
            QPushButton#SecondaryBtn:hover {
                background: rgba(255,255,255,0.82);
                border: 1px solid rgba(169, 199, 221, 0.85);
            }
            QSpinBox#TextInput {
                background: rgba(255,255,255,0.80);
                color: #2F3A40;
                border: 1px solid rgba(192, 206, 196, 0.95);
                border-radius: 12px; padding: 4px 8px;
                min-width: 64px;
            }
            QSpinBox#TextInput::up-button, QSpinBox#TextInput::down-button {
                border: none;
                width: 14px;
            }
            QSpinBox#TextInput::up-arrow, QSpinBox#TextInput::down-arrow {
                image: none;
            }
            QLabel#HintText { color: #9FB1BF; font-size: 12px; }
            QListWidget#PresetList {
                background: rgba(255,255,255,0.70);
                border: 1px solid rgba(192, 206, 196, 0.95);
                border-radius: 12px;
                color: #2F3A40;
            }
            QListWidget#PresetList::item {
                padding: 8px 10px;
                margin: 4px;
                border-radius: 10px;
                background: rgba(255,255,255,0.78);
                border: 1px solid rgba(192, 206, 196, 0.9);
            }
            QListWidget#PresetList::item:hover {
                border: 1px solid rgba(169, 199, 221, 0.8);
            }
            QScrollBar:vertical {
                background: transparent;
                width: 10px;
                margin: 6px 2px;
            }
            QScrollBar::handle:vertical {
                background: rgba(169, 199, 221, 0.6);
                border-radius: 5px;
                min-height: 24px;
            }
            QScrollBar::add-line:vertical, QScrollBar::sub-line:vertical {
                height: 0px;
            }
            """
        )

    def _apply_shadows(self, widgets):
        for widget in widgets:
            effect = QtWidgets.QGraphicsDropShadowEffect(self)
            effect.setBlurRadius(28)
            effect.setXOffset(0)
            effect.setYOffset(10)
            effect.setColor(QtGui.QColor(0, 0, 0, 140))
            widget.setGraphicsEffect(effect)

    def _apply_window_effects(self):
        if not is_windows():
            return
        try:
            import ctypes
            from ctypes import wintypes

            hwnd = int(self.winId())

            DWMWA_WINDOW_CORNER_PREFERENCE = 33
            DWMWCP_ROUND = 2
            ctypes.windll.dwmapi.DwmSetWindowAttribute(
                wintypes.HWND(hwnd),
                ctypes.c_uint(DWMWA_WINDOW_CORNER_PREFERENCE),
                ctypes.byref(ctypes.c_int(DWMWCP_ROUND)),
                ctypes.sizeof(ctypes.c_int),
            )
            DWMWA_USE_IMMERSIVE_DARK_MODE = 20
            DWMWA_CAPTION_COLOR = 35
            DWMWA_TEXT_COLOR = 36
            dark_mode = ctypes.c_int(0)
            ctypes.windll.dwmapi.DwmSetWindowAttribute(
                wintypes.HWND(hwnd),
                ctypes.c_uint(DWMWA_USE_IMMERSIVE_DARK_MODE),
                ctypes.byref(dark_mode),
                ctypes.sizeof(dark_mode),
            )
            caption = ctypes.c_uint(0x00FFFFFF)
            ctypes.windll.dwmapi.DwmSetWindowAttribute(
                wintypes.HWND(hwnd),
                ctypes.c_uint(DWMWA_CAPTION_COLOR),
                ctypes.byref(caption),
                ctypes.sizeof(caption),
            )
            text = ctypes.c_uint(0x00000000)
            ctypes.windll.dwmapi.DwmSetWindowAttribute(
                wintypes.HWND(hwnd),
                ctypes.c_uint(DWMWA_TEXT_COLOR),
                ctypes.byref(text),
                ctypes.sizeof(text),
            )
        except Exception:
            pass

    def _init_glow(self):
        self.glow = GlowOverlay(self.centralWidget())
        self.glow.setGeometry(self.centralWidget().rect())
        app = QtWidgets.QApplication.instance()
        if app:
            app.installEventFilter(self)

    def eventFilter(self, obj, event):
        if event.type() == QtCore.QEvent.Type.MouseMove and self.centralWidget() is not None:
            pos = self.centralWidget().mapFromGlobal(event.globalPosition().toPoint())
            if self.glow:
                self.glow.set_pos(pos)
        if event.type() == QtCore.QEvent.Type.Leave:
            if self.glow:
                self.glow.set_pos(None)
        return super().eventFilter(obj, event)

    def resizeEvent(self, event):
        super().resizeEvent(event)
        if self.glow and self.centralWidget() is not None:
            self.glow.setGeometry(self.centralWidget().rect())

    def _toggle_custom_fonts(self, state):
        self.use_custom_fonts = state == QtCore.Qt.CheckState.Checked
        self.font_zh_combo.setEnabled(self.use_custom_fonts)
        self.font_en_combo.setEnabled(self.use_custom_fonts)
        self.font_combo.setEnabled(not self.use_custom_fonts)
        self.btn_save_preset.setEnabled(self.use_custom_fonts)
        self.custom_box.setVisible(self.use_custom_fonts)

    def _toggle_log(self, state):
        self.enable_log = state == QtCore.Qt.CheckState.Checked
        if self.enable_log:
            self.log_box.show()
        else:
            self.log_box.hide()

    def _toggle_custom_name(self, checked: bool):
        self.custom_name_enabled = bool(checked)
        self.name_edit.setEnabled(True)
        self.name_edit.setReadOnly(not self.custom_name_enabled)
        if self.custom_name_enabled:
            self.name_edit.setFocus()

    def _toggle_page_numbers(self, checked: bool):
        self.enable_page_numbers = bool(checked)
        if hasattr(self, "page_spin"):
            self.page_spin.setEnabled(self.enable_page_numbers)
        if hasattr(self, "page_show_spin"):
            self.page_show_spin.setEnabled(self.enable_page_numbers)

    def _pick_output_dir(self):
        directory = QtWidgets.QFileDialog.getExistingDirectory(self, "选择导出文件夹", self.output_dir or self._get_default_output_dir())
        if directory:
            self.output_dir = directory
            if hasattr(self, "output_dir_edit"):
                self.output_dir_edit.setText(directory)
            self._save_settings()

    def _commit_output_dir_text(self):
        if not hasattr(self, "output_dir_edit"):
            return
        text = self.output_dir_edit.text().strip()
        if not text:
            text = self._get_default_output_dir()
            self.output_dir_edit.setText(text)
        if os.path.isdir(text):
            self.output_dir = text
            self._save_settings()
        else:
            QtWidgets.QMessageBox.warning(self, "提示", "导出文件夹不存在，请重新选择。")
            self.output_dir_edit.setText(self.output_dir or self._get_default_output_dir())

    def _load_presets(self):
        if not os.path.isfile(self.presets_file):
            self._refresh_preset_list()
            return
        try:
            with open(self.presets_file, "r", encoding="utf-8") as f:
                data = json.load(f)
            if isinstance(data, list):
                for item in data:
                    label = item.get("label")
                    filename = item.get("filename")
                    east = item.get("east_asia")
                    latin = item.get("latin")
                    if not label or not filename:
                        continue
                    if filename not in self.font_name_map:
                        self.font_presets.append((label, filename))
                        self.font_name_map[filename] = (east or "宋体", latin or "Times New Roman")
                        self.font_combo.addItem(label, filename)
        except Exception:
            pass
        self._refresh_preset_list()

    def _refresh_preset_list(self):
        if not hasattr(self, "preset_list"):
            return
        self.preset_list.clear()
        for label, filename in self.font_presets:
            east, latin = self.font_name_map.get(filename, ("宋体", "Times New Roman"))
            item = QtWidgets.QListWidgetItem(label)
            item.setData(QtCore.Qt.ItemDataRole.UserRole, {
                "label": label,
                "filename": filename,
                "east_asia": east,
                "latin": latin,
            })
            self.preset_list.addItem(item)

    def _save_preset(self):
        if not self.use_custom_fonts:
            QtWidgets.QMessageBox.information(self, "提示", "请先勾选自定义字体方案。")
            return
        east = self.font_zh_combo.currentData()
        latin = self.font_en_combo.currentData()
        default_label = f"{east} + {latin}（自定义）"
        label, ok = QtWidgets.QInputDialog.getText(self, "保存为默认方案", "方案名称：", text=default_label)
        if not ok or not label.strip():
            return
        filename = self._create_font_template(label.strip(), east, latin)
        if not filename:
            QtWidgets.QMessageBox.warning(self, "提示", "保存失败。")
            return
        if filename not in self.font_name_map:
            self.font_presets.append((label.strip(), filename))
            self.font_name_map[filename] = (east or "宋体", latin or "Times New Roman")
            self.font_combo.addItem(label.strip(), filename)
        self._save_presets_file()
        self._refresh_preset_list()

    def _save_presets_file(self):
        items = []
        for label, filename in self.font_presets:
            if filename.endswith(".docx") and filename not in {
                "academic_songti_times.docx",
                "modern_yahei_calibri.docx",
                "official_fangsong_times.docx",
                "bold_heiti_arial.docx",
                "calligraphy_kaiti_times.docx",
            }:
                east, latin = self.font_name_map.get(filename, ("宋体", "Times New Roman"))
                items.append({
                    "label": label,
                    "filename": filename,
                    "east_asia": east,
                    "latin": latin,
                })
        try:
            with open(self.presets_file, "w", encoding="utf-8") as f:
                json.dump(items, f, ensure_ascii=False, indent=2)
        except Exception:
            pass

    def _create_font_template(self, label: str, east: str, latin: str):
        base = os.path.join(os.path.dirname(__file__), "templates")
        ref = os.path.join(base, "reference.docx")
        if not os.path.isfile(ref):
            return None
        safe = re.sub(r"[^0-9A-Za-z\u4e00-\u9fff]+", "_", label).strip("_")
        safe = safe or "custom"
        filename = f"preset_{safe}.docx"
        dest = os.path.join(base, filename)
        try:
            shutil.copyfile(ref, dest)
            patch_styles_fonts(dest, east, latin)
            patch_numbering_fonts(dest, east, latin)
            return filename
        except Exception:
            return None

    def _get_reference_doc(self):
        base = os.path.join(os.path.dirname(__file__), "templates")
        if not self.use_custom_fonts and hasattr(self, "font_combo"):
            filename = self.font_combo.currentData()
            if filename:
                chosen = os.path.join(base, filename)
                if os.path.isfile(chosen):
                    return chosen
        fallback = os.path.join(base, "reference.docx")
        if os.path.isfile(fallback):
            return fallback
        return None

    def _get_reference_fonts(self):
        if self.use_custom_fonts:
            east = self.font_zh_combo.currentData() if hasattr(self, "font_zh_combo") else None
            latin = self.font_en_combo.currentData() if hasattr(self, "font_en_combo") else None
            return (east or "宋体", latin or "Times New Roman")
        if hasattr(self, "font_combo"):
            filename = self.font_combo.currentData()
            if filename and filename in self.font_name_map:
                return self.font_name_map[filename]
        return ("宋体", "Times New Roman")

    def _log_file(self, message: str):
        try:
            with open(self.log_file, "a", encoding="utf-8") as f:
                f.write(message + "\n")
        except Exception:
            pass

    def _log(self, message: str):
        self._log_file(message)
        if self.enable_log:
            self.log_area.appendPlainText(message)

    def _log_error_file(self, message: str):
        detail = message + "\n" + traceback.format_exc()
        self._log_file(detail)

    def _log_error(self, message: str):
        detail = message + "\n" + traceback.format_exc()
        self._log(detail)

    def _set_busy(self, busy: bool):
        if busy:
            self._busy_count += 1
            if self._busy_count == 1:
                QtWidgets.QApplication.setOverrideCursor(QtCore.Qt.CursorShape.WaitCursor)
        else:
            self._busy_count = max(0, self._busy_count - 1)
            if self._busy_count == 0:
                QtWidgets.QApplication.restoreOverrideCursor()

    def _run_pandoc(self, cmd):
        kwargs = {"capture_output": True, "text": True}
        if is_windows():
            kwargs["creationflags"] = subprocess.CREATE_NO_WINDOW
            startupinfo = subprocess.STARTUPINFO()
            startupinfo.dwFlags |= subprocess.STARTF_USESHOWWINDOW
            startupinfo.wShowWindow = subprocess.SW_HIDE
            kwargs["startupinfo"] = startupinfo
        return subprocess.run(cmd, **kwargs)

    def _handle_exception(self, exc_type, exc_value, exc_traceback):
        detail = "".join(traceback.format_exception(exc_type, exc_value, exc_traceback))
        self._log("未捕获异常:\n" + detail)

    def _default_output_base(self) -> str:
        title = self.windowTitle().replace(" ", "")
        if not title:
            title = "Markdown转Word"
        return f"{title}已转换"

    def _sanitize_filename(self, name: str) -> str:
        name = re.sub(r'[\\\\/:*?"<>|]', "_", name).strip()
        return name

    def _unique_output_path(self, directory: str, base: str) -> str:
        candidate = os.path.join(directory, base + ".docx")
        if not os.path.exists(candidate):
            return candidate
        idx = 2
        while True:
            candidate = os.path.join(directory, f"{base}_{idx}.docx")
            if not os.path.exists(candidate):
                return candidate
            idx += 1

    def _get_output_path(self, directory: str, base_override: str | None = None) -> str:
        if self.custom_name_enabled:
            base = self.name_edit.text().strip()
            if not base and base_override:
                base = base_override
        else:
            base = base_override or self._default_output_base()
        base = self._sanitize_filename(base)
        if not base:
            base = self._default_output_base()
        return self._unique_output_path(directory, base)

    def _get_default_output_dir(self) -> str:
        desktop = os.path.join(os.path.expanduser("~"), "Desktop")
        if os.path.isdir(desktop):
            return desktop
        return os.getcwd()

    def _resolve_output_dir(self, fallback_dir: str | None = None) -> str:
        candidate = ""
        if hasattr(self, "output_dir_edit"):
            candidate = self.output_dir_edit.text().strip()
        if not candidate:
            candidate = self.output_dir
        if candidate and os.path.isdir(candidate):
            return candidate
        if fallback_dir and os.path.isdir(fallback_dir):
            return fallback_dir
        return self._get_default_output_dir()

    def _format_pandoc_error(self, stderr: str) -> str:
        stderr = (stderr or "").strip()
        lower = stderr.lower()
        tips = []
        if "permission denied" in lower:
            tips.append("目标文件被占用或无权限。请关闭 Word/WPS，或更换文件名。")
        if "note with key" in lower:
            tips.append("脚注/注释编号重复。请检查类似 [^1] 的标记，或删除该引用。")
        if "yaml" in lower and "metadata" in lower:
            tips.append("YAML 元数据格式错误。请移除或修正文档开头的 YAML 区块。")
        if not tips:
            tips.append("请检查输入内容与输出路径。")
        msg = "转换失败。\n\n解决方法：\n" + "\n".join(f"- {t}" for t in tips)
        if stderr:
            msg += "\n\n原始信息：\n" + stderr
        return msg

    def _digest_text(self, text: str) -> str:
        return hashlib.sha1(text.encode("utf-8")).hexdigest()

    def _split_blocks_for_ai(self, text: str) -> list[str]:
        lines = text.splitlines()
        blocks = []
        in_code = False
        table_mode = False
        current_table = []
        current_code = []

        def flush_table():
            nonlocal current_table, table_mode
            if current_table:
                blocks.append(current_table)
                current_table = []
            table_mode = False

        def flush_code():
            nonlocal current_code, in_code
            if current_code:
                blocks.append(current_code)
                current_code = []
            in_code = False

        def is_table_separator(line: str) -> bool:
            return bool(re.match(r"^\|?[\s:-]+\|[\s:-]*$", line))

        i = 0
        while i < len(lines):
            line = lines[i]
            stripped = line.strip()

            if stripped.startswith("```"):
                if in_code:
                    current_code.append(line)
                    flush_code()
                else:
                    flush_table()
                    in_code = True
                    current_code.append(line)
                i += 1
                continue

            if in_code:
                current_code.append(line)
                i += 1
                continue

            if not stripped:
                flush_table()
                i += 1
                continue

            if table_mode:
                if "|" in line or is_table_separator(stripped):
                    current_table.append(line)
                    i += 1
                    continue
                flush_table()
                continue

            next_line = lines[i + 1] if i + 1 < len(lines) else ""
            if "|" in line and is_table_separator(next_line.strip()):
                flush_table()
                table_mode = True
                current_table.append(line)
                current_table.append(next_line)
                i += 2
                continue

            blocks.append([line])
            i += 1

        flush_table()
        flush_code()

        out = []
        for block in blocks:
            content = "\n".join(block).rstrip()
            if content.strip():
                out.append(content)
        return out

    def _build_ai_blocks(self, text: str) -> list[dict]:
        blocks = []
        for idx, content in enumerate(self._split_blocks_for_ai(text), start=1):
            digest = hashlib.sha1(content.encode("utf-8")).hexdigest()[:6]
            anchor = ""
            for line in content.splitlines():
                line = line.strip()
                if line:
                    anchor = line
                    break
            if len(anchor) > 40:
                anchor = anchor[:40]
            anchor = anchor.replace('"', "'")
            blocks.append(
                {
                    "id": f"B{idx:03d}",
                    "hash": digest,
                    "anchor": anchor,
                    "text": content,
                }
            )
        return blocks

    def _build_ai_prompt(self, blocks: list[dict]) -> str:
        lines = [
            "你是论文格式检查器，只允许返回 JSON，不得改写正文。",
            "你只能标注结构或补齐摘要/关键词/结论。",
            "",
            "【可用操作】",
            "- heading(level)",
            "- section(type=abstract|keywords|conclusion|references|appendix)",
            "- insert_section(type=abstract|keywords|conclusion)",
            "- ignore",
            "",
            "【返回规则】",
            "- 只返回需要修改的 block",
            "- 必须包含 block_id 和 hash",
            "- insert_section 必须包含 title + content + after_block",
            "- content 字数：摘要 200~350，关键词 3~6 个，结论 120~220",
            "- 不确定的放入 warnings",
            "",
            "【Blocks】",
        ]
        for block in blocks:
            lines.append(f"[{block['id']}] anchor=\"{block['anchor']}\" hash={block['hash']}")
        return "\n".join(lines)

    def _reset_ai_task(self):
        self.ai_task_text = None
        self.ai_task_blocks = []
        self.ai_task_digest = None
        self.ai_task_mode = None
        self.ai_task_output_dir = None
        self.ai_task_base_override = None

    def _save_ai_task_state(self, text: str, blocks: list[dict], mode: str, output_dir: str, base_override: str | None):
        self.ai_task_text = text
        self.ai_task_blocks = blocks
        self.ai_task_digest = self._digest_text(text)
        self.ai_task_mode = mode
        self.ai_task_output_dir = output_dir
        self.ai_task_base_override = base_override

    def _export_ai_task_docx(self, text: str, output_path: str, blocks: list[dict]) -> bool:
        try:
            doc = Document()
            doc.add_heading("修改任务", level=0)
            doc.add_heading("原文", level=1)
            for line in text.splitlines():
                doc.add_paragraph(line)
            doc.add_heading("Blocks（含完整内容）", level=1)
            for block in blocks:
                anchor = block.get("anchor", "").replace('"', "'")
                doc.add_paragraph(f"[{block['id']}] anchor=\"{anchor}\" hash={block['hash']}")
                doc.add_paragraph("-----")
                for line in block.get("text", "").splitlines():
                    doc.add_paragraph(line)
                doc.add_paragraph("-----")
            doc.save(output_path)
            return True
        except Exception:
            self._log_error("生成修改任务失败")
            return False

    def _read_text_from_file(self, path: str) -> str | None:
        ext = os.path.splitext(path)[1].lower()
        if ext in {".md", ".txt"}:
            return read_text_file(path)
        if ext == ".docx":
            if not self.pandoc:
                QtWidgets.QMessageBox.critical(self, "缺少 Pandoc", "未找到 pandoc。")
                return None
            cmd = [self.pandoc, "-f", "docx", "-t", "markdown", path]
            self._log("运行: " + " ".join(cmd))
            result = subprocess.run(cmd, capture_output=True, text=True)
            if result.returncode != 0:
                msg = self._format_pandoc_error(result.stderr)
                QtWidgets.QMessageBox.critical(self, "转换失败", msg)
                self._log(result.stderr.strip() or "Pandoc 执行失败")
                return None
            return result.stdout
        return None

    def _get_task_output_path(self, directory: str) -> str:
        base = "修改任务"
        return self._unique_output_path(directory, base)

    def _copy_ai_prompt(self):
        raw = None
        mode = self.source_mode
        if mode == "text":
            raw = self.text_edit.toPlainText().strip()
        elif self.ai_task_text:
            raw = self.ai_task_text
        if not raw:
            QtWidgets.QMessageBox.warning(self, "提示", "Markdown 内容为空。")
            return
        blocks = self._build_ai_blocks(raw)
        if not blocks:
            QtWidgets.QMessageBox.warning(self, "提示", "未识别到可用文本块。")
            return
        prompt = self._build_ai_prompt(blocks)
        clipboard = QtWidgets.QApplication.clipboard()
        clipboard.setText(prompt)
        if mode == "text":
            self._save_ai_task_state(raw, blocks, "text", self._get_default_output_dir(), self._default_output_base())
        self._log("已复制 AI 提示词到剪贴板")
        QtWidgets.QMessageBox.information(self, "完成", "AI 提示词已复制到剪贴板。")

    def _apply_ai_from_clipboard(self):
        if not self.ai_task_digest or not self.ai_task_text:
            QtWidgets.QMessageBox.warning(self, "提示", "请先生成修改任务并让 AI 返回 JSON。")
            return
        raw_text = self.ai_task_text
        if self.ai_task_mode == "text":
            current = self.text_edit.toPlainText().strip()
            if current and current != raw_text:
                QtWidgets.QMessageBox.warning(self, "提示", "文本已变化，请重新生成修改任务。")
                return
        raw_json = ""
        if hasattr(self, "ai_json_edit"):
            raw_json = self.ai_json_edit.toPlainText().strip()
        if not raw_json:
            clipboard = QtWidgets.QApplication.clipboard()
            raw_json = clipboard.text().strip()
        try:
            payload = json.loads(raw_json)
        except Exception:
            QtWidgets.QMessageBox.warning(self, "提示", "JSON 格式无效，请检查输入。")
            return
        actions = payload.get("actions", [])
        warnings = payload.get("warnings", [])
        if not isinstance(actions, list):
            QtWidgets.QMessageBox.warning(self, "提示", "JSON 格式错误：actions 必须是数组。")
            return
        blocks = self.ai_task_blocks or self._build_ai_blocks(raw_text)
        block_map = {block["id"]: block for block in blocks}
        allowed_ops = {"heading", "section", "insert_section", "ignore"}
        section_titles = {
            "abstract": "摘要",
            "keywords": "关键词",
            "conclusion": "结论",
            "references": "参考文献",
            "appendix": "附录",
        }
        block_actions = {}
        insert_after = {}
        applied = 0
        skipped = 0
        for action in actions:
            if not isinstance(action, dict):
                skipped += 1
                continue
            op = action.get("op")
            if op not in allowed_ops:
                skipped += 1
                continue
            if op == "insert_section":
                sec_type = action.get("type")
                content = (action.get("content") or "").strip()
                after_block = action.get("after_block")
                title = action.get("title") or section_titles.get(sec_type, "")
                if not title or not content:
                    skipped += 1
                    continue
                block_text = f"# {title}\n{content}".strip()
                key = after_block if after_block in block_map else "__end__"
                insert_after.setdefault(key, []).append(block_text)
                applied += 1
                continue
            block_id = action.get("block_id")
            block_hash = action.get("hash")
            block = block_map.get(block_id)
            if not block or not block_hash or block_hash != block["hash"]:
                skipped += 1
                continue
            if op == "ignore":
                continue
            confidence = float(action.get("confidence", 1.0))
            current = block_actions.get(block_id)
            if current and confidence <= current["confidence"]:
                continue
            block_actions[block_id] = {"action": action, "confidence": confidence}

        def apply_heading(text: str, level: int) -> str:
            lines = text.splitlines()
            if not lines:
                return text
            level = max(1, min(level, 6))
            title = re.sub(r"^#+\\s*", "", lines[0]).strip()
            lines[0] = f"{'#' * level} {title}".strip()
            return "\n".join(lines)

        def apply_section(text: str, sec_type: str) -> str:
            title = section_titles.get(sec_type, "").strip()
            if not title:
                return text
            heading = f"# {title}"
            lines = text.splitlines()
            if not lines:
                return heading
            first = lines[0].strip()
            if first.startswith("#"):
                lines[0] = heading
                return "\n".join(lines)
            if first == title or first.startswith(title + "：") or first.startswith(title + ":"):
                rest = first[len(title):].lstrip("：:").strip()
                lines[0] = heading
                if rest:
                    lines.insert(1, rest)
                return "\n".join(lines)
            return f"{heading}\n{text}".strip()

        new_blocks = []
        for block in blocks:
            text = block["text"]
            action = block_actions.get(block["id"])
            if action:
                op = action["action"].get("op")
                if op == "heading":
                    level = int(action["action"].get("level", 1))
                    text = apply_heading(text, level)
                elif op == "section":
                    sec_type = action["action"].get("type")
                    text = apply_section(text, sec_type)
            new_blocks.append(text)
            extra = insert_after.get(block["id"])
            if extra:
                new_blocks.extend(extra)
        if "__end__" in insert_after:
            new_blocks.extend(insert_after["__end__"])

        new_text = "\n\n".join([block for block in new_blocks if block.strip()])
        if new_text:
            if self.ai_task_mode == "text":
                self.text_edit.setPlainText(new_text)
                self._log(f"AI 应用完成：{applied} 条，忽略 {skipped} 条。")
                if isinstance(warnings, list) and warnings:
                    self._log(f"AI 警告：{len(warnings)} 条")
                QtWidgets.QMessageBox.information(self, "完成", "AI 结果已应用到文本。")
            elif self.ai_task_mode == "file":
                output_dir = self.ai_task_output_dir or self._get_default_output_dir()
                output_path = self._get_output_path(output_dir, self.ai_task_base_override)
                self._convert_markdown(new_text, output_path)
                self._log(f"AI 应用完成：{applied} 条，忽略 {skipped} 条。")
                if isinstance(warnings, list) and warnings:
                    self._log(f"AI 警告：{len(warnings)} 条")
            else:
                self.text_edit.setPlainText(new_text)
                QtWidgets.QMessageBox.information(self, "完成", "AI 结果已应用到文本。")
        else:
            QtWidgets.QMessageBox.warning(self, "提示", "应用后内容为空，已取消。")

    def _preprocess_text_preview(self, text: str) -> str:
        text = sanitize_special_marks(text)
        text = merge_vertical_text(text)
        text = normalize_fenced_math(text)
        text = normalize_math_blocks(text)
        text = strip_standalone_brackets(text)
        text = filter_horizontal_rules(text, self.keep_rules)
        text = normalize_ai_headings(text)
        return text

    def _update_text_preview(self):
        if self.source_mode != "text":
            return
        raw = self.text_edit.toPlainText()
        preview = self._preprocess_text_preview(raw)
        if preview != self.output_text.toPlainText():
            self.output_text.setPlainText(preview)

    def _set_source_mode(self, mode: str):
        self.source_mode = mode
        self.btn_file.setChecked(mode == "file")
        self.btn_text.setChecked(mode == "text")
        self.stack.setCurrentIndex(0 if mode == "file" else 1)
        self.btn_export_text.setVisible(mode == "text")
        self.right_stack.setCurrentIndex(0)
        if mode == "text":
            self._update_text_preview()

    def _handle_drop(self, paths):
        for path in paths:
            self._convert_file(path)

    def _pick_files(self):
        files, _ = QtWidgets.QFileDialog.getOpenFileNames(
            self, "选择文件", "",
            "Markdown (*.md);;Text (*.txt);;All Files (*.*)"
        )
        for path in files:
            self._convert_file(path)

    def _convert_file(self, path: str):
        if not path:
            return
        ext = os.path.splitext(path)[1].lower()
        if ext not in {".md", ".txt"}:
            QtWidgets.QMessageBox.warning(self, "提示", "仅支持 .md / .txt 文件。")
            return
        base = os.path.splitext(os.path.basename(path))[0] + "已转换"
        out_dir = self._resolve_output_dir(os.path.dirname(path))
        output_path = self._get_output_path(out_dir, base)
        markdown_text = read_text_file(path)
        self._convert_markdown(markdown_text, output_path)

    def _generate_ai_task_for_file(self, path: str):
        if not path:
            return
        ext = os.path.splitext(path)[1].lower()
        if ext not in {".md", ".txt", ".docx"}:
            QtWidgets.QMessageBox.warning(self, "提示", "仅支持 .md / .txt / .docx 文件。")
            return
        text = self._read_text_from_file(path)
        if text is None:
            return
        blocks = self._build_ai_blocks(text)
        if not blocks:
            QtWidgets.QMessageBox.warning(self, "提示", "未识别到可用文本块。")
            return
        output_path = self._get_task_output_path(os.path.dirname(path))
        if not self._export_ai_task_docx(text, output_path, blocks):
            QtWidgets.QMessageBox.warning(self, "提示", "修改任务生成失败，已记录日志。")
            return
        base_override = os.path.splitext(os.path.basename(path))[0] + "已转换"
        self._save_ai_task_state(text, blocks, "file", os.path.dirname(path), base_override)
        self._on_export_success(output_path)
        QtWidgets.QMessageBox.information(self, "完成", "已生成修改任务，请将该文档发给 AI。")

    def _convert_docx(self, input_path: str, output_path: str, list_widget=None):
        reference = self._get_reference_doc()
        font_pair = self._get_reference_fonts()
        enable_outline = self.enable_heading_outline
        keep_rules = self.keep_rules
        page_enable = self.enable_page_numbers
        page_start = self.page_start
        page_start_page = self.page_start_page
        bold_levels = set(self.bold_levels) if hasattr(self, "bold_levels") else {1, 2}
        def task():
            return self._convert_docx_task(
                input_path, output_path, reference, font_pair, keep_rules, enable_outline, bold_levels,
                page_enable, page_start, page_start_page
            )
        self._run_conversion(task, list_widget)

    def _convert_markdown(self, markdown_text: str, output_path: str, list_widget=None):
        reference = self._get_reference_doc()
        font_pair = self._get_reference_fonts()
        keep_rules = self.keep_rules
        enable_outline = self.enable_heading_outline
        page_enable = self.enable_page_numbers
        page_start = self.page_start
        page_start_page = self.page_start_page
        bold_levels = set(self.bold_levels) if hasattr(self, "bold_levels") else {1, 2}
        def task():
            return self._convert_markdown_task(
                markdown_text, output_path, reference, font_pair, keep_rules, enable_outline, bold_levels,
                page_enable, page_start, page_start_page
            )
        self._run_conversion(task, list_widget)

    def _run_conversion(self, task_func, list_widget=None):
        self._set_busy(True)
        worker = BackgroundTask(task_func)
        worker.signals.finished.connect(lambda result: self._handle_task_finished(result, list_widget))
        worker.signals.error.connect(self._handle_task_error)
        self.thread_pool.start(worker)

    def _handle_task_finished(self, result, list_widget=None):
        self._set_busy(False)
        if not result:
            return
        if not result.get("ok"):
            msg = result.get("message", "转换失败。")
            msg = f"{msg}\n\n日志位置：{self.log_file}"
            QtWidgets.QMessageBox.critical(self, result.get("title", "转换失败"), msg)
            return
        warnings = result.get("warnings") or []
        if warnings:
            QtWidgets.QMessageBox.warning(self, "提示", "\n".join(warnings))
        self._on_export_success(result["output_path"], list_widget)

    def _handle_task_error(self, title: str, message: str):
        self._set_busy(False)
        msg = f"{message}\n\n日志位置：{self.log_file}"
        QtWidgets.QMessageBox.critical(self, title, msg)

    def _convert_docx_task(self, input_path: str, output_path: str, reference: str | None, font_pair,
                           keep_rules: bool, enable_outline: bool, bold_levels: set,
                           page_enable: bool, page_start: int, page_start_page: int):
        if not self.pandoc:
            return {"ok": False, "title": "缺少 Pandoc", "message": "未找到 pandoc。"}
        tmp_md = None
        try:
            with tempfile.NamedTemporaryFile(delete=False, suffix=".md") as tmp:
                tmp_md = tmp.name
            md_cmd = [self.pandoc, "-f", "docx", "-t", "markdown", "-o", tmp_md, input_path]
            self._log_file("运行: " + " ".join(md_cmd))
            md_result = self._run_pandoc(md_cmd)
            if md_result.returncode != 0:
                msg = self._format_pandoc_error(md_result.stderr)
                self._log_file(md_result.stderr.strip() or "Pandoc 执行失败")
                return {"ok": False, "title": "转换失败", "message": msg}
            markdown_text = read_text_file(tmp_md).strip()
        finally:
            if tmp_md:
                try:
                    os.remove(tmp_md)
                except Exception:
                    pass
        if not markdown_text:
            return {"ok": False, "title": "转换失败", "message": "Word 文档内容为空或无法读取。"}
        markdown_text = unescape_markdown_from_docx(markdown_text)
        return self._convert_markdown_task(
            markdown_text, output_path, reference, font_pair, keep_rules, enable_outline, bold_levels,
            page_enable, page_start, page_start_page
        )

    def _convert_markdown_task(self, markdown_text: str, output_path: str, reference: str | None,
                               font_pair, keep_rules: bool, enable_outline: bool, bold_levels: set,
                               page_enable: bool, page_start: int, page_start_page: int):
        if not self.pandoc:
            return {"ok": False, "title": "缺少 Pandoc", "message": "未找到 pandoc。"}

        markdown_text = sanitize_special_marks(markdown_text)
        markdown_text = ensure_closed_code_blocks(markdown_text)
        markdown_text = merge_vertical_text(markdown_text)
        markdown_text = normalize_fenced_math(markdown_text)
        markdown_text = normalize_table_breaks(markdown_text)
        markdown_text = normalize_markdown_tables(markdown_text)
        markdown_text = normalize_math_blocks(markdown_text)
        markdown_text = strip_standalone_brackets(markdown_text)
        markdown_text = normalize_inline_math_parens(markdown_text)
        markdown_text = filter_horizontal_rules(markdown_text, keep_rules)
        markdown_text = normalize_ai_headings(markdown_text)
        with tempfile.NamedTemporaryFile(delete=False, suffix=".md", mode="w", encoding="utf-8") as tmp:
            tmp.write(markdown_text)
            tmp_path = tmp.name

        try:
            cmd = [
                self.pandoc,
                "-f",
                "markdown+tex_math_dollars+tex_math_single_backslash+tex_math_double_backslash+raw_tex-yaml_metadata_block+autolink_bare_uris",
                "-t",
                "docx",
                "-o",
                output_path,
            ]
            if reference:
                cmd.extend(["--reference-doc", reference])
            cmd.append(tmp_path)
            self._log_file("运行: " + " ".join(cmd))
            result = self._run_pandoc(cmd)
            if result.returncode != 0:
                msg = self._format_pandoc_error(result.stderr)
                self._log_file(result.stderr.strip() or "Pandoc 执行失败")
                return {"ok": False, "title": "转换失败", "message": msg}
        finally:
            try:
                os.remove(tmp_path)
            except Exception:
                pass

        warnings = []
        try:
            east, latin = font_pair
            patch_numbering_fonts(output_path, east, latin)
            patch_styles_fonts(output_path, east, latin)
            patch_heading_styles(output_path, enable_outline, None)
            patch_document_outline(output_path, enable_outline)
            if not enable_outline:
                demote_heading_paragraphs(output_path, bold_levels=bold_levels)
            remove_literal_math_markers(output_path)
            apply_table_grid_style(output_path)
            update_page_numbers(output_path, page_enable, page_start, page_start_page)
        except Exception:
            self._log_error_file("编号字体修复失败")
        return {"ok": True, "output_path": output_path, "warnings": warnings}

    def _export_from_text(self):
        if self.source_mode != "text":
            QtWidgets.QMessageBox.warning(self, "提示", "请切换到文本导入后再导出。")
            return
        markdown_text = self.text_edit.toPlainText().strip()
        if not markdown_text:
            QtWidgets.QMessageBox.warning(self, "提示", "Markdown 内容为空。")
            return
        if self.ai_enabled:
            blocks = self._build_ai_blocks(markdown_text)
            if not blocks:
                QtWidgets.QMessageBox.warning(self, "提示", "未识别到可用文本块。")
                return
            output_dir = self._get_default_output_dir()
            output_path = self._get_task_output_path(output_dir)
            if not self._export_ai_task_docx(markdown_text, output_path, blocks):
                QtWidgets.QMessageBox.warning(self, "提示", "修改任务生成失败，已记录日志。")
                return
            self._save_ai_task_state(markdown_text, blocks, "text", output_dir, self._default_output_base())
            self._on_export_success(output_path)
            QtWidgets.QMessageBox.information(self, "完成", "已生成修改任务，请将该文档发给 AI。")
            return
        output_dir = self._resolve_output_dir(None)
        path = self._get_output_path(output_dir)
        self._convert_markdown(markdown_text, path)

    def _on_export_success(self, output_path: str, list_widget=None):
        if list_widget is None:
            list_widget = self.export_list
        list_widget.add_item(output_path)
        self._add_history(output_path)
        if self.open_after_export:
            try:
                os.startfile(output_path)
            except Exception:
                pass
        self._log(f"已保存：{output_path}")

    def _load_settings(self):
        default_dir = self._get_default_output_dir()
        self.output_dir = default_dir
        if hasattr(self, "output_dir_edit"):
            self.output_dir_edit.setText(default_dir)
        if not os.path.isfile(self.settings_file):
            return
        try:
            with open(self.settings_file, "r", encoding="utf-8") as f:
                data = json.load(f)
            if isinstance(data, dict):
                saved_dir = data.get("output_dir") or ""
                if saved_dir and os.path.isdir(saved_dir):
                    self.output_dir = saved_dir
                    if hasattr(self, "output_dir_edit"):
                        self.output_dir_edit.setText(saved_dir)
        except Exception:
            pass

    def _save_settings(self):
        try:
            payload = {"output_dir": self.output_dir}
            with open(self.settings_file, "w", encoding="utf-8") as f:
                json.dump(payload, f, ensure_ascii=False, indent=2)
        except Exception:
            pass

    def _load_export_history(self):
        if not os.path.isfile(self.history_file):
            return
        try:
            with open(self.history_file, "r", encoding="utf-8") as f:
                data = json.load(f)
            if not isinstance(data, list):
                return
            self._export_history = []
            for item in data:
                if not isinstance(item, str):
                    continue
                path = os.path.normpath(item)
                if os.path.isfile(path):
                    self._export_history.append(path)
            for path in self._export_history:
                self.export_list.add_item(path)
        except Exception:
            pass

    def _save_export_history(self):
        try:
            with open(self.history_file, "w", encoding="utf-8") as f:
                json.dump(self._export_history, f, ensure_ascii=False, indent=2)
        except Exception:
            pass

    def _add_history(self, path: str):
        if not path:
            return
        clean = os.path.normpath(path)
        key = os.path.normcase(clean)
        self._export_history = [p for p in self._export_history if os.path.normcase(p) != key]
        self._export_history.insert(0, clean)
        self._export_history = self._export_history[:200]
        self._save_export_history()

    def _remove_history(self, path: str):
        if not path:
            return
        key = os.path.normcase(os.path.normpath(path))
        self._export_history = [p for p in self._export_history if os.path.normcase(p) != key]
        self._save_export_history()


def main():
    app = QtWidgets.QApplication(sys.argv)
    icon_path = os.path.join(os.path.dirname(__file__), "assets", "app.ico")
    app_icon = QtGui.QIcon(icon_path) if os.path.exists(icon_path) else QtGui.QIcon()
    if not app_icon.isNull():
        app.setWindowIcon(app_icon)
    window = MainWindow()
    if not app_icon.isNull():
        window.setWindowIcon(app_icon)
    window.show()
    window._apply_window_effects()
    sys.exit(app.exec())


if __name__ == "__main__":
    main()


